import sqlite3
import pandas as pd
import numpy as np
import csv
import io
import json
import requests
import os
import glob
from pathlib import Path
from functools import wraps
from datetime import datetime, date, timedelta
from zoneinfo import ZoneInfo

from flask import (
    Blueprint,
    jsonify,
    render_template,
    session,
    redirect,
    url_for,
    request,
    send_file,
    current_app
)
from services.quiz_generator import generate_quiz
from sklearn.cluster import KMeans
from sklearn.preprocessing import StandardScaler

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    PageBreak
)
from docx import Document
from modules.ai_report_generator import (
    check_ollama,
    generate_candidate_report,
    generate_exam_report
)
from modules.datasciencemodule import (
    calculate_face_presence,
    calculate_integrity_score,
    build_score_distribution,
    build_risk_distribution,
    build_violation_heatmap,
    perform_behavioral_clustering,
    build_cohort_analysis
)
from services.quiz_generator import generate_quiz
# ============================================================
# BLUEPRINT
# ============================================================

admin_dashboard_bp = Blueprint(
    "admin_dashboard",
    __name__,
    url_prefix="/admin"
)


# ============================================================
# DATABASE
# ============================================================

BASE_DIR = Path(__file__).resolve().parent.parent

DATABASE = BASE_DIR / "database.db"


def get_db_connection():

    conn = sqlite3.connect(DATABASE)

    conn.row_factory = sqlite3.Row

    conn.execute(
        "PRAGMA foreign_keys = ON"
    )

    return conn

# ============================================================
# OLLAMA CONFIGURATION
# ============================================================

OLLAMA_URL = os.getenv(
    "OLLAMA_URL",
    "http://127.0.0.1:11434/api/generate"
)

OLLAMA_MODEL = os.getenv(
    "OLLAMA_MODEL",
    "llama3.2:3b"
)

# ============================================================
# ADMIN AUTHENTICATION DECORATORS
# ============================================================

def admin_login_required(function):
    """Decorator for admin HTML page routes: redirects unauthenticated requests to /admin/login."""
    @wraps(function)
    def wrapper(*args, **kwargs):
        if "admin_id" not in session:
            return redirect(url_for("admin.admin_login_page"))
        return function(*args, **kwargs)
    return wrapper


def admin_api_required(function):
    """Decorator for admin JSON API routes: returns JSON 401 if unauthenticated."""
    @wraps(function)
    def wrapper(*args, **kwargs):
        if "admin_id" not in session:
            return jsonify({
                "success": False,
                "status": "error",
                "message": "Unauthorized access. Admin authentication required."
            }), 401
        return function(*args, **kwargs)
    return wrapper


# Backward-compatible alias
admin_required = admin_api_required

# ============================================================
# ADMIN LOGOUT PAGE ROUTE
# ============================================================

@admin_dashboard_bp.route("/logout")
def admin_logout_page():
    session.clear()
    return redirect(url_for("admin.admin_login_page"))

@admin_dashboard_bp.route("/support-requests")
@admin_login_required
def support_requests_page():
    """Admin support requests page."""
    return render_template("support_request.html")


# ============================================================
# ADMIN DASHBOARD PAGE
# ============================================================
@admin_dashboard_bp.route("/dashboard")
@admin_login_required
def dashboard():

    print("Inside dashboard")
    print(dict(session))

    return render_template("admin_dashboard.html")
# ============================================================
# EXAMINATIONS PAGE
# ============================================================

@admin_dashboard_bp.route("/examinations")
@admin_login_required
def examinations_page():
    return render_template("examinations.html")


# ==========================================================
# AI GENERATED REPORTS PAGE
# ==========================================================

@admin_dashboard_bp.route(
    "/ai-generated-reports",
    methods=["GET"]
)
@admin_login_required
def ai_generated_reports_page():
    return render_template("ai_generated_reports.html")


@admin_dashboard_bp.route(
    "/api/support/<int:ticket_id>",
    methods=["PUT"]
)
def update_support_ticket(ticket_id):

    if "admin_id" not in session:
        return jsonify({
            "success": False,
            "message": "Admin not logged in"
        }), 401

    admin_id = session["admin_id"]

    data = request.get_json(silent=True)

    if not data:
        return jsonify({
            "success": False,
            "message": "Invalid request data"
        }), 400

    status = str(
        data.get("status", "")
    ).strip()

    admin_response = str(
        data.get("admin_response", "")
    ).strip()

    allowed_statuses = [
        "Open",
        "In Progress",
        "Resolved",
        "Closed"
    ]

    if status not in allowed_statuses:
        return jsonify({
            "success": False,
            "message": "Invalid ticket status"
        }), 400

    conn = get_db_connection()

    try:

        ticket = conn.execute("""
            SELECT id
            FROM SupportTickets
            WHERE id = ?
        """, (ticket_id,)).fetchone()

        if ticket is None:
            return jsonify({
                "success": False,
                "message": "Support ticket not found"
            }), 404

        resolved_at = None

        if status in ["Resolved", "Closed"]:
            resolved_at = datetime.now().isoformat()

        conn.execute("""
            UPDATE SupportTickets

            SET
                status = ?,
                admin_response = ?,
                assigned_admin_id = ?,
                updated_at = CURRENT_TIMESTAMP,
                resolved_at = ?

            WHERE id = ?
        """, (
            status,
            admin_response,
            admin_id,
            resolved_at,
            ticket_id
        ))

        conn.commit()

        return jsonify({
            "success": True,
            "message": "Support ticket updated successfully"
        })

    finally:
        conn.close()
@admin_dashboard_bp.route("/api/support", methods=["GET"])
@admin_login_required
def get_admin_support_tickets():

    if "admin_id" not in session:
        return jsonify({
            "success": False,
            "message": "Admin not logged in"
        }), 401

    conn = get_db_connection()

    try:

        tickets = conn.execute("""
            SELECT
                s.id,

                s.candidate_id,

                c.name AS candidate_name,

                c.email AS candidate_email,

                s.issue_type,

                s.priority,

                s.subject,

                s.message,

                s.status,

                s.admin_response,

                s.created_at,

                s.updated_at,

                s.resolved_at,

                a.full_name AS admin_name

            FROM SupportTickets s

            INNER JOIN Candidates c
                ON c.id = s.candidate_id

            LEFT JOIN Admins a
                ON a.id = s.assigned_admin_id

            ORDER BY
                CASE s.priority
                    WHEN 'Critical' THEN 1
                    WHEN 'High' THEN 2
                    WHEN 'Medium' THEN 3
                    WHEN 'Low' THEN 4
                    ELSE 5
                END,

                datetime(s.created_at) DESC
        """).fetchall()

        return jsonify({
            "success": True,
            "tickets": [
                dict(ticket)
                for ticket in tickets
            ]
        })

    finally:
        conn.close()


@admin_dashboard_bp.route("/profile")
@admin_api_required
def profile_page():
    return render_template("admin_profile.html")

# ============================================================
# ADMIN PROFILE API
# ============================================================
@admin_dashboard_bp.route("/api/admin-profile", methods=["GET"])
@admin_dashboard_bp.route("/api/dashboard/admin-profile", methods=["GET"])
@admin_dashboard_bp.route("/api/profile", methods=["GET"])
@admin_api_required
def admin_profile():

    conn = get_db_connection()

    try:

        # Get logged-in admin ID from session
        admin_id = session.get("admin_id")

        if not admin_id:
            return jsonify({
                "success": False,
                "message": "Admin session expired."
            }), 401


        # Fetch admin details
        admin = conn.execute(
            """
            SELECT
                id,
                full_name,
                email,
                employee_id,
                username,
                created_at
            FROM Admins
            WHERE id = ?
            """,
            (admin_id,)
        ).fetchone()


        if admin is None:
            return jsonify({
                "success": False,
                "message": "Admin not found."
            }), 404


        return jsonify({

            "success": True,

            "admin": {

                "id": admin["id"],

                "name": admin["full_name"] or "",

                "full_name": admin["full_name"] or "",

                "email": admin["email"] or "",

                "employee_id": admin["employee_id"] or "",

                "username": admin["username"] or "",

                "created_at": admin["created_at"] or ""

            }

        }), 200


    except Exception as error:

        import traceback

        print("Admin profile error:", repr(error))

        traceback.print_exc()

        return jsonify({

            "success": False,

            "message": "Unable to load admin profile.",

            "error": str(error)

        }), 500


    finally:

        conn.close()


# ============================================================
# STATUS OF EACH CANDIDATE PAGE
# ============================================================

@admin_dashboard_bp.route(
    "/status-of-candidates",
    methods=["GET"]
)
@admin_login_required
def candidate_status_page():
    return render_template(
        "status_of_candidates.html"
    )



# ============================================================
# CANDIDATE STATUS API
# ============================================================
@admin_dashboard_bp.route(
    "/api/candidate-status",
    methods=["GET"]
)
@admin_dashboard_bp.route(
    "/api/candidate_status",
    methods=["GET"]
)
@admin_api_required
def candidate_status_api():

    conn = get_db_connection()

    try:

        exam_id = request.args.get(
            "exam_id",
            type=int
        )

        # ============================================================
        # EXAMS
        # ============================================================

        exams = conn.execute(
            """
            SELECT
                id,
                title,
                topic,
                duration,
                start_time,
                end_time
            FROM Exams
            ORDER BY
                datetime(start_time) DESC
            """
        ).fetchall()

        # ============================================================
        # LATEST SESSION FOR EACH CANDIDATE
        # ============================================================

        if exam_id is not None:

            session_condition = "s.exam_id = ?"
            session_params = (exam_id,)

        else:

            session_condition = "1 = 1"
            session_params = ()

        query = f"""
            WITH latest_sessions AS (

                SELECT
                    s.*,

                    ROW_NUMBER() OVER (
                        PARTITION BY s.candidate_id
                        ORDER BY
                            datetime(
                                COALESCE(
                                    s.login_time,
                                    '1970-01-01'
                                )
                            ) DESC,
                            s.id DESC
                    ) AS rn

                FROM SessionLogs s

                WHERE {session_condition}
            )

            SELECT

                c.id AS candidate_id,
                c.name,
                c.email,
                c.photo_path,

                ls.id AS session_id,
                ls.exam_id,
                ls.login_time,
                ls.logout_time,
                ls.status AS session_status,

                e.title AS exam_title,
                e.topic AS exam_topic,

                (
                    SELECT COUNT(*)
                    FROM ViolationLogs v
                    WHERE
                        v.candidate_id = c.id
                        AND (
                            ls.exam_id IS NULL
                            OR v.exam_id = ls.exam_id
                        )
                ) AS violation_count,

                (
                    SELECT COUNT(*)
                    FROM ViolationLogs v
                    WHERE
                        v.candidate_id = c.id
                        AND (
                            ls.exam_id IS NULL
                            OR v.exam_id = ls.exam_id
                        )
                        AND v.evidence_image IS NOT NULL
                ) AS evidence_count,

                (
                    SELECT MAX(v.violation_time)
                    FROM ViolationLogs v
                    WHERE
                        v.candidate_id = c.id
                        AND (
                            ls.exam_id IS NULL
                            OR v.exam_id = ls.exam_id
                        )
                ) AS last_violation

            FROM Candidates c

            LEFT JOIN latest_sessions ls
                ON ls.candidate_id = c.id
                AND ls.rn = 1

            LEFT JOIN Exams e
                ON e.id = ls.exam_id

            ORDER BY
                c.name COLLATE NOCASE ASC
        """

        rows = conn.execute(
            query,
            session_params
        ).fetchall()

        # ============================================================
        # BUILD CANDIDATES
        # ============================================================

        candidates = []

        for row in rows:

            candidate_id = row["candidate_id"]

            # ========================================================
            # VIOLATIONS
            # ========================================================

            violation_count = int(
                row["violation_count"] or 0
            )

            evidence_count = int(
                row["evidence_count"] or 0
            )

            # ========================================================
            # SESSION STATUS
            # ========================================================

            raw_status = (
                row["session_status"]
                or ""
            ).strip().lower()

            logout_time = row["logout_time"]

            if (
                raw_status in (
                    "active",
                    "online",
                    "running"
                )
                and not logout_time
            ):

                current_status = "Online"

            elif violation_count > 0:

                current_status = "Warning"

            else:

                current_status = "Offline"

            # ========================================================
            # RISK LEVEL
            # ========================================================

            if violation_count >= 4:

                risk_level = "High"

            elif violation_count >= 2:

                risk_level = "Medium"

            else:

                risk_level = "Low"

            # ========================================================
            # INTEGRITY SCORE
            # ========================================================

            integrity_score = max(
                0,
                min(
                    100,
                    100 - (
                        violation_count * 10
                    )
                )
            )

            # ========================================================
            # LAST ACTIVITY
            # ========================================================

            last_activity = (
                row["last_violation"]
                or row["login_time"]
            )

            # ========================================================
            # COMPLETED EXAMS
            # ========================================================

            completed_exam_rows = conn.execute(
                """
                SELECT
                    ea.id AS attempt_id,
                    ea.exam_id,
                    ea.score,

                    e.title AS exam_title,
                    e.topic AS exam_topic

                FROM ExamAttempts ea

                LEFT JOIN Exams e
                    ON e.id = ea.exam_id

                WHERE
                    ea.candidate_id = ?

                ORDER BY
                    ea.id DESC
                """,
                (candidate_id,)
            ).fetchall()

            # ========================================================
            # BUILD EXAM HISTORY
            # ========================================================

            completed_exams = []

            for attempt in completed_exam_rows:

                score = attempt["score"]

                try:

                    score = float(
                        score or 0
                    )

                except (
                    TypeError,
                    ValueError
                ):

                    score = 0

                completed_exams.append({

                    "attempt_id":
                        attempt["attempt_id"],

                    "exam_id":
                        attempt["exam_id"],

                    "exam_name":
                        attempt["exam_title"]
                        or "Unknown Exam",

                    "topic":
                        attempt["exam_topic"]
                        or "",

                    "score":
                        score

                })

            # ========================================================
            # REMOVE DUPLICATE EXAMS
            # ========================================================
            #
            # ExamAttempts is ordered by id DESC.
            #
            # Therefore, the first record for an exam is
            # the latest attempt.
            #
            # ========================================================

            unique_exams = {}

            for exam in completed_exams:

                exam_key = exam["exam_id"]

                if exam_key is None:
                    continue

                if exam_key not in unique_exams:

                    unique_exams[
                        exam_key
                    ] = exam

            completed_exams = list(
                unique_exams.values()
            )

            # ========================================================
            # EXAM COUNT
            # ========================================================

            exams_completed = len(
                completed_exams
            )

            # ========================================================
            # AVERAGE EXAM SCORE
            # ========================================================

            if completed_exams:

                average_exam_score = round(
                    sum(
                        exam["score"]
                        for exam in completed_exams
                    )
                    /
                    len(completed_exams),
                    1
                )

            else:

                average_exam_score = 0

            # ========================================================
            # CURRENT EXAM SCORE
            # ========================================================

            current_exam_score = None

            current_exam_id = row["exam_id"]

            if current_exam_id is not None:

                for exam in completed_exams:

                    if (
                        exam["exam_id"]
                        ==
                        current_exam_id
                    ):

                        current_exam_score = (
                            exam["score"]
                        )

                        break

            # ========================================================
            # CANDIDATE OBJECT
            # ========================================================

            candidates.append({

                # ----------------------------------------------------
                # BASIC INFORMATION
                # ----------------------------------------------------

                "candidate_id":
                    candidate_id,

                "name":
                    row["name"]
                    or "Unknown",

                "email":
                    row["email"]
                    or "",

                "photo":
                    row["photo_path"]
                    or "",

                # ----------------------------------------------------
                # SESSION INFORMATION
                # ----------------------------------------------------

                "session_id":
                    row["session_id"],

                "exam_id":
                    row["exam_id"],

                "exam_title":
                    row["exam_title"]
                    or "No examination",

                "exam_topic":
                    row["exam_topic"]
                    or "",

                "login_time":
                    row["login_time"],

                "logout_time":
                    row["logout_time"],

                # ----------------------------------------------------
                # STATUS
                # ----------------------------------------------------

                "status":
                    current_status,

                # ----------------------------------------------------
                # RISK
                # ----------------------------------------------------

                "risk":
                    risk_level,

                # ----------------------------------------------------
                # INTEGRITY
                # ----------------------------------------------------

                "violation_count":
                    violation_count,

                "evidence_count":
                    evidence_count,

                "integrity_score":
                    integrity_score,

                "last_activity":
                    last_activity,

                # ----------------------------------------------------
                # EXAM INFORMATION
                # ----------------------------------------------------

                "exams_completed":
                    exams_completed,

                "completed_exams":
                    completed_exams,

                "average_exam_score":
                    average_exam_score,

                "current_exam_score":
                    current_exam_score

            })

        # ============================================================
        # STATISTICS
        # ============================================================

        total = len(candidates)

        online = sum(
            1
            for candidate in candidates
            if candidate["status"] == "Online"
        )

        warning = sum(
            1
            for candidate in candidates
            if candidate["status"] == "Warning"
        )

        violations = sum(
            1
            for candidate in candidates
            if candidate["violation_count"] > 0
        )

        offline = sum(
            1
            for candidate in candidates
            if candidate["status"] == "Offline"
        )

        # ============================================================
        # ACTIVITY FEED
        # ============================================================

        activity_query = """
            SELECT

                v.id,
                v.candidate_id,
                v.exam_id,
                v.violation_type,
                v.face_count,
                v.evidence_image,
                v.violation_time,

                c.name AS candidate_name,

                e.title AS exam_title

            FROM ViolationLogs v

            LEFT JOIN Candidates c
                ON c.id = v.candidate_id

            LEFT JOIN Exams e
                ON e.id = v.exam_id

            WHERE
                (
                    ? IS NULL
                    OR v.exam_id = ?
                )

            ORDER BY
                datetime(v.violation_time) DESC,
                v.id DESC

            LIMIT 20
        """

        activity_rows = conn.execute(
            activity_query,
            (
                exam_id,
                exam_id
            )
        ).fetchall()

        activity_feed = []

        for row in activity_rows:

            violation_type = (
                row["violation_type"]
                or "Integrity event"
            )

            severity = (
                "High"
                if violation_type.upper()
                in (
                    "MULTIPLE_FACES",
                    "UNKNOWN_FACE",
                    "IDENTITY MISMATCH",
                    "IDENTITY MISMATCH."
                )
                else "Medium"
            )

            activity_feed.append({

                "id":
                    row["id"],

                "candidate_id":
                    row["candidate_id"],

                "candidate_name":
                    row["candidate_name"]
                    or "Unknown Candidate",

                "exam_id":
                    row["exam_id"],

                "exam_title":
                    row["exam_title"]
                    or "Unknown Examination",

                "type":
                    violation_type,

                "violation_type":
                    violation_type,

                "face_count":
                    row["face_count"]
                    or 0,

                "evidence_image":
                    row["evidence_image"],

                "time":
                    row["violation_time"],

                "severity":
                    severity

            })

        # ============================================================
        # REALTIME ACTIVITY
        # ============================================================

        chart_query = """
            SELECT

                strftime(
                    '%H:%M',
                    violation_time
                ) AS time_label,

                COUNT(*) AS event_count

            FROM ViolationLogs

            WHERE
                violation_time IS NOT NULL

                AND (
                    ? IS NULL
                    OR exam_id = ?
                )

            GROUP BY
                strftime(
                    '%H:%M',
                    violation_time
                )

            ORDER BY
                datetime(
                    MAX(violation_time)
                ) ASC

            LIMIT 12
        """

        chart_rows = conn.execute(
            chart_query,
            (
                exam_id,
                exam_id
            )
        ).fetchall()

        realtime_activity = [

            {
                "time":
                    row["time_label"],

                "count":
                    int(
                        row["event_count"] or 0
                    )
            }

            for row in chart_rows

        ]

        # ============================================================
        # FINAL RESPONSE
        # ============================================================

        return jsonify({

            "success": True,

            # --------------------------------------------------------
            # STATISTICS
            # --------------------------------------------------------

            "statistics": {

                "total":
                    total,

                "online":
                    online,

                "warning":
                    warning,

                "violations":
                    violations,

                "offline":
                    offline

            },

            # --------------------------------------------------------
            # EXAMS
            # --------------------------------------------------------

            "exams": [

                {

                    "id":
                        exam["id"],

                    "title":
                        exam["title"],

                    "topic":
                        exam["topic"],

                    "duration":
                        exam["duration"],

                    "start_time":
                        exam["start_time"],

                    "end_time":
                        exam["end_time"]

                }

                for exam in exams

            ],

            # --------------------------------------------------------
            # CANDIDATES
            # --------------------------------------------------------

            "candidates":
                candidates,

            # --------------------------------------------------------
            # ACTIVITY
            # --------------------------------------------------------

            "activity":
                activity_feed,

            "activity_feed":
                activity_feed,

            # --------------------------------------------------------
            # REALTIME ACTIVITY
            # --------------------------------------------------------

            "realtime_activity":
                realtime_activity

        })

    except Exception as error:

        print(
            "Candidate status error:",
            repr(error)
        )

        return jsonify({

            "success":
                False,

            "message":
                "Unable to load candidate status.",

            "error":
                str(error)

        }), 500

    finally:

        conn.close()


# ============================================================
# DASHBOARD STATISTICS
# ============================================================

@admin_dashboard_bp.route(
    "/api/dashboard-stats"
)
@admin_dashboard_bp.route(
    "/api/dashboard/stats"
)
@admin_api_required
def dashboard_stats():

    conn = get_db_connection()

    try:

        # ----------------------------------------------------
        # TOTAL CANDIDATES
        # ----------------------------------------------------

        total_candidates = conn.execute(
            """
            SELECT COUNT(*)
            FROM Candidates
            """
        ).fetchone()[0]


        # ----------------------------------------------------
        # TOTAL EXAMS
        # ----------------------------------------------------

        total_exams = conn.execute(
            """
            SELECT COUNT(*)
            FROM Exams
            """
        ).fetchone()[0]


        # ----------------------------------------------------
        # RUNNING EXAMS
        # ----------------------------------------------------

        running_exams = conn.execute(
            """
            SELECT COUNT(*)
            FROM Exams
            WHERE datetime(start_time)
                  <= datetime('now','localtime')

              AND datetime(end_time)
                  >= datetime('now','localtime')
            """
        ).fetchone()[0]


        # ----------------------------------------------------
        # COMPLETED TODAY
        # ----------------------------------------------------

        completed_today = conn.execute(
            """
            SELECT COUNT(*)
            FROM ExamAttempts
            WHERE date(submitted_at)
                  = date('now','localtime')
            """
        ).fetchone()[0]


        # ----------------------------------------------------
        # TOTAL VIOLATIONS
        # ----------------------------------------------------

        violation_count = conn.execute(
            """
            SELECT COUNT(*)
            FROM ViolationLogs
            """
        ).fetchone()[0]


        # ----------------------------------------------------
        # HIGH RISK
        # ----------------------------------------------------

        high_risk = conn.execute(
            """
            SELECT COUNT(*)
            FROM IntegrityScores
            WHERE UPPER(risk_label) = 'HIGH'
            """
        ).fetchone()[0]


        # ----------------------------------------------------
        # REPORT COUNT
        # ----------------------------------------------------

        report_count = conn.execute(
            """
            SELECT COUNT(*)
            FROM AIReports
            """
        ).fetchone()[0]


        # ----------------------------------------------------
        # AVERAGE INTEGRITY
        # ----------------------------------------------------

        average_integrity = conn.execute(
            """
            SELECT AVG(integrity_score)
            FROM IntegrityScores
            """
        ).fetchone()[0]


        if average_integrity is None:

            average_integrity = 0

        else:

            average_integrity = round(
                average_integrity,
                1
            )


        return jsonify({

            "success": True,

            "statistics": {

                "total_candidates":
                    total_candidates,

                "total_exams":
                    total_exams,

                "running_exams":
                    running_exams,

                "completed_today":
                    completed_today,

                "violation_count":
                    violation_count,

                "high_risk":
                    high_risk,

                "report_count":
                    report_count,

                "average_integrity":
                    average_integrity

            }

        })


    except Exception as error:

        print(
            "Dashboard statistics error:",
            error
        )

        return jsonify({

            "success": False,

            "message":
                "Unable to load dashboard statistics",

            "error":
                str(error)

        }), 500


    finally:

        conn.close()


# ============================================================
# INTEGRITY SUMMARY
# ============================================================

@admin_dashboard_bp.route(
    "/api/dashboard/integrity"
)
@admin_required
def integrity_summary():

    conn = get_db_connection()

    try:

        # Average integrity

        average = conn.execute(
            """
            SELECT AVG(integrity_score)
            FROM IntegrityScores
            """
        ).fetchone()[0]


        # Face presence

        face_ratio = conn.execute(
            """
            SELECT AVG(face_presence_ratio)
            FROM IntegrityScores
            """
        ).fetchone()[0]


        # Warnings

        warnings = conn.execute(
            """
            SELECT COALESCE(
                SUM(warning_count),
                0
            )
            FROM IntegrityScores
            """
        ).fetchone()[0]


        # Risk counts

        high = conn.execute(
            """
            SELECT COUNT(*)
            FROM IntegrityScores
            WHERE UPPER(risk_label) = 'HIGH'
            """
        ).fetchone()[0]


        medium = conn.execute(
            """
            SELECT COUNT(*)
            FROM IntegrityScores
            WHERE UPPER(risk_label) = 'MEDIUM'
            """
        ).fetchone()[0]


        low = conn.execute(
            """
            SELECT COUNT(*)
            FROM IntegrityScores
            WHERE UPPER(risk_label) = 'LOW'
            """
        ).fetchone()[0]


        average = (
            round(average, 1)
            if average is not None
            else 0
        )


        face_percentage = (
            round(face_ratio * 100, 1)
            if face_ratio is not None
            else 0
        )


        return jsonify({

            "success": True,

            "overallIntegrity":
                average,

            "averageIntegrity":
                average,

            "facePresenceRatio":
                face_percentage,

            "warningCount":
                warnings,

            "highRisk":
                high,

            "mediumRisk":
                medium,

            "lowRisk":
                low,

            "integritySummary": [

                "Face detection monitoring active",

                "Identity verification available",

                "Tab switching monitoring active",

                "Screen and environment monitoring active"

            ]

        })


    except Exception as error:

        print(
            "Integrity summary error:",
            error
        )

        return jsonify({

            "success": False,

            "message":
                "Unable to load integrity summary",

            "error":
                str(error)

        }), 500


    finally:

        conn.close()


# ============================================================
# RISK DISTRIBUTION
# ============================================================

@admin_dashboard_bp.route(
    "/api/dashboard/risk-distribution"
)
@admin_required
def risk_distribution():

    conn = get_db_connection()

    try:

        rows = conn.execute(
            """
            SELECT
                UPPER(risk_label) AS risk,
                COUNT(*) AS count
            FROM IntegrityScores
            GROUP BY UPPER(risk_label)
            """
        ).fetchall()


        distribution = {

            "LOW": 0,

            "MEDIUM": 0,

            "HIGH": 0

        }


        for row in rows:

            risk = row["risk"]

            if risk in distribution:

                distribution[risk] = (
                    row["count"]
                )


        total = sum(
            distribution.values()
        )


        result = []


        for risk, count in distribution.items():

            percentage = (

                round(
                    count / total * 100,
                    1
                )

                if total > 0

                else 0

            )


            result.append({

                "risk":
                    risk,

                "count":
                    count,

                "percentage":
                    percentage

            })


        return jsonify({

            "success": True,

            "distribution":
                result,

            "total":
                total

        })


    except Exception as error:

        print(
            "Risk distribution error:",
            error
        )

        return jsonify({

            "success": False,

            "message":
                "Unable to load risk distribution",

            "error":
                str(error)

        }), 500


    finally:

        conn.close()


# ============================================================
# INTEGRITY TREND
# ============================================================

@admin_dashboard_bp.route(
    "/api/dashboard/integrity-trend"
)
@admin_required
def integrity_trend():

    conn = get_db_connection()

    try:

        rows = conn.execute(
            """
            SELECT
                date(generated_at) AS day,
                AVG(integrity_score) AS score
            FROM IntegrityScores
            WHERE generated_at >= datetime(
                'now',
                '-6 days',
                'localtime'
            )
            GROUP BY date(generated_at)
            ORDER BY day ASC
            """
        ).fetchall()


        data = {}


        today = date.today()


        for i in range(7):

            current = (
                today -
                timedelta(days=6 - i)
            )


            data[
                current.strftime(
                    "%Y-%m-%d"
                )
            ] = 0


        for row in rows:

            if row["day"] in data:

                data[
                    row["day"]
                ] = round(
                    row["score"],
                    1
                )


        labels = []

        values = []


        for day, score in data.items():

            parsed = datetime.strptime(
                day,
                "%Y-%m-%d"
            )


            labels.append(
                parsed.strftime("%a")
            )

            values.append(score)


        return jsonify({

            "success": True,

            "labels":
                labels,

            "values":
                values

        })


    except Exception as error:

        print(
            "Integrity trend error:",
            error
        )

        return jsonify({

            "success": False,

            "message":
                "Unable to load integrity trend",

            "error":
                str(error)

        }), 500


    finally:

        conn.close()


# ============================================================
# LIVE EXAMS
# ============================================================

@admin_dashboard_bp.route(
    "/api/dashboard/live-exams"
)
@admin_required
def live_exams():

    conn = get_db_connection()

    try:

        rows = conn.execute(
            """
            SELECT
                e.id,
                e.title,
                e.topic,
                e.duration,
                e.total_questions,
                e.start_time,
                e.end_time,

                (
                    SELECT COUNT(*)
                    FROM SessionLogs s
                    WHERE s.status = 'active'
                ) AS active_candidates,

                (
                    SELECT COUNT(*)
                    FROM ViolationLogs v
                    WHERE v.exam_id = e.id
                ) AS violations

            FROM Exams e

            WHERE datetime(e.start_time)
                  <= datetime('now','localtime')

              AND datetime(e.end_time)
                  >= datetime('now','localtime')

            ORDER BY
                datetime(e.start_time) ASC
            """
        ).fetchall()


        exams = []


        for row in rows:

            progress = calculate_exam_progress(
                row["start_time"],
                row["end_time"]
            )


            exams.append({

                "id":
                    row["id"],

                "examName":
                    row["title"],

                "title":
                    row["title"],

                "status":
                    "Running",

                "candidates":
                    row["active_candidates"] or 0,

                "violations":
                    row["violations"] or 0,

                "progress":
                    progress,

                "timeLeft":
                    calculate_time_left(
                        row["end_time"]
                    ),

                "startTime":
                    row["start_time"],

                "endTime":
                    row["end_time"]

            })


        return jsonify({

            "success": True,

            "exams":
                exams,

            "liveExamTable":
                exams

        })


    except Exception as error:

        print(
            "Live exams error:",
            error
        )

        return jsonify({

            "success": False,

            "message":
                "Unable to load live examinations",

            "error":
                str(error)

        }), 500


    finally:

        conn.close()


# ============================================================
# DATETIME PARSER
# ============================================================

def parse_datetime(value):

    if not value:

        raise ValueError(
            "Invalid datetime"
        )


    value = str(value).strip()


    formats = [

        "%Y-%m-%d %H:%M:%S",

        "%Y-%m-%d %H:%M",

        "%Y-%m-%dT%H:%M:%S",

        "%Y-%m-%dT%H:%M",

        "%d-%m-%Y %H:%M:%S",

        "%d-%m-%Y %H:%M"

    ]


    for fmt in formats:

        try:

            return datetime.strptime(
                value,
                fmt
            )

        except ValueError:

            continue


    raise ValueError(
        f"Unsupported datetime format: {value}"
    )


# ============================================================
# EXAM PROGRESS
# ============================================================

def calculate_exam_progress(
    start_time,
    end_time
):

    try:

        start = parse_datetime(
            start_time
        )

        end = parse_datetime(
            end_time
        )

        now = datetime.now()


        total = (
            end - start
        ).total_seconds()


        elapsed = (
            now - start
        ).total_seconds()


        if total <= 0:

            return 0


        progress = (
            elapsed / total
        ) * 100


        return max(
            0,
            min(
                100,
                round(progress)
            )
        )


    except Exception:

        return 0


# ============================================================
# TIME LEFT
# ============================================================

def calculate_time_left(end_time):

    try:

        end = parse_datetime(
            end_time
        )

        now = datetime.now()


        seconds = int(
            (
                end - now
            ).total_seconds()
        )


        if seconds <= 0:

            return "00:00:00"


        hours = (
            seconds // 3600
        )


        minutes = (
            seconds % 3600
        ) // 60


        remaining = (
            seconds % 60
        )


        return (
            f"{hours:02d}:"
            f"{minutes:02d}:"
            f"{remaining:02d}"
        )


    except Exception:

        return "--:--:--"


# ============================================================
# RECENT ACTIVITY
# ============================================================

@admin_dashboard_bp.route(
    "/api/dashboard/activity"
)
@admin_required
def recent_activity():

    conn = get_db_connection()

    try:

        activities = []


        # ----------------------------------------------------
        # VIOLATIONS
        # ----------------------------------------------------

        violations = conn.execute(
            """
            SELECT
                v.violation_time,
                v.violation_type,
                c.name AS candidate_name,
                e.title AS exam_title
            FROM ViolationLogs v

            INNER JOIN Candidates c
                ON c.id = v.candidate_id

            INNER JOIN Exams e
                ON e.id = v.exam_id

            ORDER BY
                datetime(v.violation_time) DESC

            LIMIT 10
            """
        ).fetchall()


        for row in violations:

            activities.append({

                "time":
                    row["violation_time"],

                "type":
                    "Violation",

                "title":
                    "Integrity Violation",

                "description":
                    (
                        f'{row["candidate_name"]} - '
                        f'{row["violation_type"]}'
                    ),

                "message":
                    (
                        f'{row["candidate_name"]} - '
                        f'{row["violation_type"]}'
                    ),

                "exam":
                    row["exam_title"]

            })


        # ----------------------------------------------------
        # EXAM ATTEMPTS
        # ----------------------------------------------------

        attempts = conn.execute(
            """
            SELECT
                a.submitted_at,
                a.percentage,
                a.result,
                c.name AS candidate_name,
                e.title AS exam_title

            FROM ExamAttempts a

            INNER JOIN Candidates c
                ON c.id = a.candidate_id

            INNER JOIN Exams e
                ON e.id = a.exam_id

            ORDER BY
                datetime(a.submitted_at) DESC

            LIMIT 10
            """
        ).fetchall()


        for row in attempts:

            activities.append({

                "time":
                    row["submitted_at"],

                "type":
                    "Exam",

                "title":
                    "Exam Completed",

                "description":
                    (
                        f'{row["candidate_name"]} '
                        f'completed '
                        f'{row["exam_title"]}'
                    ),

                "message":
                    (
                        f'{row["candidate_name"]} '
                        f'completed '
                        f'{row["exam_title"]}'
                    ),

                "exam":
                    row["exam_title"],

                "percentage":
                    row["percentage"],

                "result":
                    row["result"]

            })


        # ----------------------------------------------------
        # LOGIN SESSIONS
        # ----------------------------------------------------

        sessions = conn.execute(
            """
            SELECT
                s.login_time,
                s.status,
                c.name AS candidate_name

            FROM SessionLogs s

            INNER JOIN Candidates c
                ON c.id = s.candidate_id

            WHERE s.login_time IS NOT NULL

            ORDER BY
                datetime(s.login_time) DESC

            LIMIT 10
            """
        ).fetchall()


        for row in sessions:

            activities.append({

                "time":
                    row["login_time"],

                "type":
                    "Login",

                "title":
                    "Candidate Login",

                "description":
                    (
                        f'{row["candidate_name"]} '
                        f'logged into an examination'
                    ),

                "message":
                    (
                        f'{row["candidate_name"]} '
                        f'logged into an examination'
                    ),

                "status":
                    row["status"]

            })


        activities.sort(
            key=lambda item:
                str(item.get("time", "")),
            reverse=True
        )


        activities = activities[:15]


        return jsonify({

            "success": True,

            "activities":
                activities,

            "recentActivity":
                activities

        })


    except Exception as error:

        print(
            "Recent activity error:",
            error
        )

        return jsonify({

            "success": False,

            "message":
                "Unable to load recent activity",

            "error":
                str(error)

        }), 500


    finally:

        conn.close()


# ============================================================
# UPCOMING EXAMS
# ============================================================

@admin_dashboard_bp.route(
    "/api/dashboard/upcoming-exams"
)
@admin_required
def upcoming_exams():

    conn = get_db_connection()

    try:

        rows = conn.execute(
            """
            SELECT
                id,
                title,
                topic,
                difficulty,
                duration,
                total_questions,
                total_marks,
                start_time,
                end_time

            FROM Exams

            WHERE datetime(start_time)
                  > datetime('now','localtime')

            ORDER BY
                datetime(start_time) ASC

            LIMIT 10
            """
        ).fetchall()


        exams = []


        for row in rows:

            exams.append({

                "id":
                    row["id"],

                "title":
                    row["title"],

                "topic":
                    row["topic"] or "",

                "difficulty":
                    row["difficulty"] or "",

                "duration":
                    row["duration"],

                "total_questions":
                    row["total_questions"],

                "total_marks":
                    row["total_marks"],

                "start_time":
                    row["start_time"],

                "end_time":
                    row["end_time"],

                # CamelCase versions too

                "totalQuestions":
                    row["total_questions"],

                "totalMarks":
                    row["total_marks"],

                "startTime":
                    row["start_time"],

                "endTime":
                    row["end_time"]

            })


        return jsonify({

            "success": True,

            "exams":
                exams,

            "upcomingExams":
                exams

        })


    except Exception as error:

        print(
            "Upcoming exams error:",
            error
        )

        return jsonify({

            "success": False,

            "message":
                "Unable to load upcoming exams",

            "error":
                str(error)

        }), 500


    finally:

        conn.close()


# ============================================================
# RECENT CANDIDATES
# ============================================================

@admin_dashboard_bp.route(
    "/api/dashboard/candidates"
)
@admin_required
def recent_candidates():

    conn = get_db_connection()

    try:

        rows = conn.execute(
            """
            SELECT
                c.id,
                c.name,
                c.email,
                c.photo_path,
                c.created_at

            FROM Candidates c

            ORDER BY
                datetime(c.created_at) DESC

            LIMIT 20
            """
        ).fetchall()


        candidates = []


        for row in rows:

            candidates.append({

                "id":
                    row["id"],

                "name":
                    row["name"],

                "email":
                    row["email"],

                "photo":
                    row["photo_path"],

                "photoPath":
                    row["photo_path"],

                "createdAt":
                    row["created_at"]

            })


        return jsonify({

            "success": True,

            "candidates":
                candidates,

            "recentCandidates":
                candidates

        })


    except Exception as error:

        print(
            "Recent candidates error:",
            error
        )

        return jsonify({

            "success": False,

            "message":
                "Unable to load candidates",

            "error":
                str(error)

        }), 500


    finally:

        conn.close()


# ============================================================
# LIVE ALERTS
# ============================================================

@admin_dashboard_bp.route(
    "/api/dashboard/alerts"
)
@admin_required
def live_alerts():

    conn = get_db_connection()

    try:

        rows = conn.execute(
            """
            SELECT
                v.id,
                v.violation_type,
                v.face_count,
                v.violation_time,

                c.name AS candidate_name,

                c.photo_path,

                e.title AS exam_title

            FROM ViolationLogs v

            INNER JOIN Candidates c
                ON c.id = v.candidate_id

            INNER JOIN Exams e
                ON e.id = v.exam_id

            ORDER BY
                datetime(v.violation_time) DESC

            LIMIT 10
            """
        ).fetchall()


        alerts = []


        for row in rows:

            violation = (
                row["violation_type"]
                or "Unknown Violation"
            )


            lower = violation.lower()


            if (
                "multiple" in lower
                or "face" in lower
                or "identity" in lower
            ):

                risk = "HIGH"


            elif (
                "tab" in lower
                or "focus" in lower
                or "screen" in lower
            ):

                risk = "MEDIUM"


            else:

                risk = "LOW"


            alerts.append({

                "id":
                    row["id"],

                "candidate":
                    row["candidate_name"],

                "candidateName":
                    row["candidate_name"],

                "exam":
                    row["exam_title"],

                "violation":
                    violation,

                "violationType":
                    violation,

                "faceCount":
                    row["face_count"],

                "risk":
                    risk,

                "time":
                    row["violation_time"],

                "photo":
                    row["photo_path"]

            })


        return jsonify({

            "success": True,

            "alerts":
                alerts,

            "liveAlerts":
                alerts

        })


    except Exception as error:

        print(
            "Live alerts error:",
            error
        )

        return jsonify({

            "success": False,

            "message":
                "Unable to load alerts",

            "error":
                str(error)

        }), 500


    finally:

        conn.close()


# ============================================================
# VIOLATIONS SUMMARY (DASHBOARD)
# ============================================================

@admin_dashboard_bp.route(
    "/api/violations",
    methods=["GET"]
)
@admin_api_required
def violations():

    conn = get_db_connection()

    try:

        rows = conn.execute(
            """
            SELECT
                v.id,
                v.violation_type,
                v.evidence_image,
                v.face_count,
                v.violation_time,

                c.name AS candidate_name,

                e.title AS exam_title

            FROM ViolationLogs v

            INNER JOIN Candidates c
                ON c.id = v.candidate_id

            INNER JOIN Exams e
                ON e.id = v.exam_id

            ORDER BY
                datetime(v.violation_time) DESC

            LIMIT 30
            """
        ).fetchall()


        result = []


        for row in rows:

            result.append({

                "id":
                    row["id"],

                "candidate":
                    row["candidate_name"],

                "exam":
                    row["exam_title"],

                "type":
                    row["violation_type"],

                "image":
                    row["evidence_image"],

                "face_count":
                    row["face_count"],

                "time":
                    row["violation_time"]

            })


        return jsonify({

            "success": True,

            "violations":
                result

        })


    except Exception as error:

        print(
            "Violations error:",
            error
        )

        return jsonify({

            "success": False,

            "message":
                "Unable to load violations",

            "error":
                str(error)

        }), 500


    finally:

        conn.close()


# ============================================================
# AI REPORTS
# ============================================================

@admin_dashboard_bp.route(
    "/api/reports"
)
@admin_required
def reports():

    conn = get_db_connection()

    try:

        rows = conn.execute(
            """
            SELECT
                i.id,
                i.candidate_id,
                i.exam_id,
                i.integrity_score,
                i.face_presence_ratio,
                i.warning_count,
                i.risk_label,
                i.generated_at,

                c.name AS candidate_name,

                e.title AS exam_title

            FROM IntegrityScores i

            INNER JOIN Candidates c
                ON c.id = i.candidate_id

            INNER JOIN Exams e
                ON e.id = i.exam_id

            ORDER BY
                datetime(i.generated_at) DESC

            LIMIT 20
            """
        ).fetchall()


        result = []


        for row in rows:

            result.append({

                "id":
                    row["id"],

                "title":
                    (
                        "Integrity Report - "
                        f'{row["candidate_name"]}'
                    ),

                "description":
                    (
                        f'{row["exam_title"]} | '
                        f'Integrity: '
                        f'{row["integrity_score"]}% | '
                        f'Risk: '
                        f'{row["risk_label"]}'
                    ),

                "candidate":
                    row["candidate_name"],

                "exam":
                    row["exam_title"],

                "integrity_score":
                    row["integrity_score"],

                "risk":
                    row["risk_label"],

                "warnings":
                    row["warning_count"],

                "generated_at":
                    row["generated_at"]

            })


        return jsonify({

            "success": True,

            "reports":
                result

        })


    except Exception as error:

        print(
            "Reports error:",
            error
        )

        return jsonify({

            "success": False,

            "message":
                "Unable to load reports",

            "error":
                str(error)

        }), 500


    finally:

        conn.close()


# ============================================================
# NOTIFICATION COUNT
# ============================================================
@admin_dashboard_bp.route(
    "/api/notifications/count"
)
@admin_login_required
def notification_count():

    conn = get_db_connection()

    try:

        count = conn.execute(
            """
            SELECT COUNT(*)
            FROM ViolationLogs

            WHERE date(violation_time)
                  = date('now','localtime')
            """
        ).fetchone()[0]

        return jsonify({

            "success": True,

            "count":
                count

        })

    except Exception as error:

        print(
            "Notification count error:",
            error
        )

        return jsonify({

            "success": False,

            "message":
                "Unable to load notification count",

            "error":
                str(error)

        }), 500

    finally:

        conn.close()


@admin_dashboard_bp.route(
    "/api/notifications",
    methods=["GET"]
)
@admin_login_required
def notifications():

    conn = get_db_connection()

    try:

        # ============================================================
        # RECENT NOTIFICATIONS
        # ============================================================

        rows = conn.execute(
            """
            SELECT

                v.id,

                v.candidate_id,

                c.name AS candidate_name,

                c.email AS candidate_email,

                v.exam_id,

                e.title AS exam_title,

                v.violation_type,

                v.face_count,

                v.violation_time

            FROM ViolationLogs v

            LEFT JOIN Candidates c
                ON c.id = v.candidate_id

            LEFT JOIN Exams e
                ON e.id = v.exam_id

            ORDER BY
                v.violation_time DESC

            LIMIT 20
            """
        ).fetchall()


        notifications = []


        for row in rows:

            violation_type = (
                row["violation_type"]
                or "Integrity Event"
            )


            notifications.append({

                "id":
                    row["id"],

                "candidateId":
                    row["candidate_id"],

                "candidateName":
                    row["candidate_name"]
                    or "Unknown Candidate",

                "candidateEmail":
                    row["candidate_email"]
                    or "",

                "examId":
                    row["exam_id"],

                "examTitle":
                    row["exam_title"]
                    or "Unknown Examination",

                "event":
                    violation_type,

                "faceCount":
                    row["face_count"]
                    or 0,

                "time":
                    row["violation_time"]
                    or "",

                "message":
                    (
                        str(violation_type)
                        + " detected for "
                        + str(
                            row["candidate_name"]
                            or "Unknown Candidate"
                        )
                    )

            })


        # ============================================================
        # TODAY'S NOTIFICATION COUNT
        # ============================================================

        count_row = conn.execute(
            """
            SELECT COUNT(*)
            FROM ViolationLogs
            WHERE date(violation_time)
                  = date('now','localtime')
            """
        ).fetchone()


        count = (
            count_row[0]
            if count_row
            else 0
        )


        return jsonify({

            "success":
                True,

            "count":
                count,

            "notifications":
                notifications

        }), 200


    except Exception as error:

        import traceback

        print(
            "Notifications API error:",
            repr(error)
        )

        traceback.print_exc()

        return jsonify({

            "success":
                False,

            "message":
                "Unable to load notifications.",

            "error":
                str(error)

        }), 500


    finally:

        conn.close()


# ============================================================
# QUIZ GENERATOR PAGE
# ============================================================

@admin_dashboard_bp.route("/quiz-generator", methods=["GET"])
@admin_dashboard_bp.route("/api/quiz-generator", methods=["GET"])
@admin_login_required
def quiz_generator_page():
    return render_template(
        "quiz_generator.html"
    )
# ============================================================
# ADMIN GENERATE QUIZ
# ============================================================

@admin_dashboard_bp.route(
    "/api/generate-quiz",
    methods=["POST"]
)
@admin_required
def admin_generate_quiz():

    try:

        data = request.get_json(silent=True) or {}

        subject = str(
            data.get("subject", "")
        ).strip()

        topic = str(
            data.get("topic", "")
        ).strip()

        difficulty = str(
            data.get("difficulty", "")
        ).strip()

        count = data.get("count")
        duration = data.get("duration")

        start_time = str(
            data.get("start_time", "")
        ).strip()

        end_time = str(
            data.get("end_time", "")
        ).strip()

        # ----------------------------------------------------
        # BASIC VALIDATION
        # ----------------------------------------------------

        if not subject:
            return jsonify({
                "status": "error",
                "message": "Subject is required."
            }), 400

        if not topic:
            return jsonify({
                "status": "error",
                "message": "Topic is required."
            }), 400

        if difficulty not in [
            "Easy",
            "Medium",
            "Hard"
        ]:
            return jsonify({
                "status": "error",
                "message": "Invalid difficulty."
            }), 400

        # ----------------------------------------------------
        # QUESTION COUNT
        # ----------------------------------------------------

        try:
            count = int(count)
        except (TypeError, ValueError):

            return jsonify({
                "status": "error",
                "message": "Invalid question count."
            }), 400

        if count < 1 or count > 150:

            return jsonify({
                "status": "error",
                "message":
                    "Question count must be between 1 and 150."
            }), 400

        # ----------------------------------------------------
        # DURATION
        # ----------------------------------------------------

        try:
            duration = int(duration)
        except (TypeError, ValueError):

            return jsonify({
                "status": "error",
                "message": "Invalid duration."
            }), 400

        if duration < 5 or duration > 300:

            return jsonify({
                "status": "error",
                "message":
                    "Duration must be between 5 and 300 minutes."
            }), 400

        # ----------------------------------------------------
        # START / END TIME
        # ----------------------------------------------------

        if not start_time:

            return jsonify({
                "status": "error",
                "message": "Start time is required."
            }), 400

        if not end_time:

            return jsonify({
                "status": "error",
                "message": "End time is required."
            }), 400

        try:

            start = datetime.fromisoformat(
                start_time
            )

            end = datetime.fromisoformat(
                end_time
            )

        except ValueError:

            return jsonify({
                "status": "error",
                "message":
                    "Invalid start or end time."
            }), 400

        if end <= start:

            return jsonify({
                "status": "error",
                "message":
                    "End time must be greater than start time."
            }), 400

        # ----------------------------------------------------
        # GENERATE QUESTIONS
        # ----------------------------------------------------

        questions = generate_quiz(
            subject,
            topic,
            difficulty,
            count
        )

        if not isinstance(questions, list):

            return jsonify({
                "status": "error",
                "message":
                    "Quiz generator returned invalid data."
            }), 500

        if len(questions) == 0:

            return jsonify({
                "status": "error",
                "message":
                    "No questions were generated."
            }), 500

        # ----------------------------------------------------
        # VALIDATE GENERATED QUESTIONS
        # ----------------------------------------------------

        required_fields = [
            "question",
            "option_a",
            "option_b",
            "option_c",
            "option_d",
            "correct_option"
        ]

        validated_questions = []

        for index, q in enumerate(
            questions,
            start=1
        ):

            if not isinstance(q, dict):

                return jsonify({
                    "status": "error",
                    "message":
                        f"Invalid question format at question {index}."
                }), 500

            for field in required_fields:

                if not str(
                    q.get(field, "")
                ).strip():

                    return jsonify({
                        "status": "error",
                        "message":
                            f"Missing {field} in question {index}."
                    }), 500

            correct_option = str(
                q["correct_option"]
            ).strip().upper()

            if correct_option not in [
                "A",
                "B",
                "C",
                "D"
            ]:

                return jsonify({
                    "status": "error",
                    "message":
                        f"Invalid correct option in question {index}."
                }), 500

            validated_questions.append({

                "question":
                    str(q["question"]).strip(),

                "option_a":
                    str(q["option_a"]).strip(),

                "option_b":
                    str(q["option_b"]).strip(),

                "option_c":
                    str(q["option_c"]).strip(),

                "option_d":
                    str(q["option_d"]).strip(),

                "correct_option":
                    correct_option

            })

        # ----------------------------------------------------
        # RESPONSE
        # ----------------------------------------------------

        return jsonify({

            "status": "success",

            "message":
                "Quiz generated successfully.",

            "questions":
                validated_questions,

            "quiz": {

                "subject":
                    subject,

                "topic":
                    topic,

                "difficulty":
                    difficulty,

                "count":
                    len(validated_questions),

                "duration":
                    duration,

                "start_time":
                    start_time,

                "end_time":
                    end_time
            }

        }), 200

    except Exception as error:

        print(
            "Admin generate quiz error:",
            repr(error)
        )

        return jsonify({

            "status":
                "error",

            "message":
                "Unable to generate quiz.",

            "error":
                str(error)

        }), 500

# ============================================================
# ADMIN SAVE QUIZ
# ============================================================

@admin_dashboard_bp.route(
    "/api/save-quiz",
    methods=["POST"]
)
@admin_api_required
def admin_save_quiz():
    conn = None
    try:

        data = request.get_json() or {}


        subject = str(
            data.get("subject", "")
        ).strip()

        topic = str(
            data.get("topic", "")
        ).strip()

        difficulty = str(
            data.get("difficulty", "")
        ).strip()

        duration = data.get("duration")

        start_time = data.get(
            "start_time",
            ""
        )

        end_time = data.get(
            "end_time",
            ""
        )

        questions = data.get(
            "questions",
            []
        )


        # ----------------------------------------------------
        # VALIDATION
        # ----------------------------------------------------

        if not subject:

            return jsonify({
                "status": "error",
                "message": "Subject is required."
            }), 400


        if not topic:

            return jsonify({
                "status": "error",
                "message": "Topic is required."
            }), 400


        if difficulty not in [
            "Easy",
            "Medium",
            "Hard"
        ]:

            return jsonify({
                "status": "error",
                "message": "Invalid difficulty."
            }), 400


        try:

            duration = int(duration)

        except (TypeError, ValueError):

            return jsonify({
                "status": "error",
                "message": "Invalid duration."
            }), 400


        if duration < 5 or duration > 300:

            return jsonify({
                "status": "error",
                "message":
                    "Duration must be between 5 and 300 minutes."
            }), 400


        if not start_time:

            return jsonify({
                "status": "error",
                "message": "Start time is required."
            }), 400


        if not end_time:

            return jsonify({
                "status": "error",
                "message": "End time is required."
            }), 400


        if not isinstance(
            questions,
            list
        ) or len(questions) == 0:

            return jsonify({
                "status": "error",
                "message":
                    "No questions available to save."
            }), 400


        # ----------------------------------------------------
        # DATE VALIDATION
        # ----------------------------------------------------

        try:

            start = datetime.fromisoformat(
                start_time
            )

            end = datetime.fromisoformat(
                end_time
            )

        except ValueError:

            return jsonify({
                "status": "error",
                "message":
                    "Invalid start or end time."
            }), 400


        if end <= start:

            return jsonify({
                "status": "error",
                "message":
                    "End time must be greater than start time."
            }), 400


        # ----------------------------------------------------
        # DATABASE
        # ----------------------------------------------------

        conn = get_db_connection()

        cursor = conn.cursor()


        title = subject

        description = (
            f"{difficulty} Level "
            f"AI Generated Quiz"
        )

        total_questions = len(
            questions
        )

        total_marks = total_questions


        # ----------------------------------------------------
        # INSERT EXAM
        # ----------------------------------------------------

        cursor.execute(
            """
            INSERT INTO Exams
            (
                title,
                topic,
                difficulty,
                description,
                duration,
                total_questions,
                total_marks,
                start_time,
                end_time
            )
            VALUES
            (
                ?, ?, ?, ?, ?, ?, ?, ?, ?
            )
            """,
            (
                title,
                topic,
                difficulty,
                description,
                duration,
                total_questions,
                total_marks,
                start_time,
                end_time
            )
        )


        exam_id = cursor.lastrowid


        # ----------------------------------------------------
        # INSERT QUESTIONS
        # ----------------------------------------------------

        for q in questions:

            if not isinstance(q, dict):

                raise ValueError(
                    "Invalid question format."
                )


            question = q.get(
                "question",
                ""
            )

            option_a = q.get(
                "option_a",
                ""
            )

            option_b = q.get(
                "option_b",
                ""
            )

            option_c = q.get(
                "option_c",
                ""
            )

            option_d = q.get(
                "option_d",
                ""
            )

            correct_option = q.get(
                "correct_option",
                ""
            )


            if not question:

                raise ValueError(
                    "Question text is missing."
                )


            if not option_a:

                raise ValueError(
                    "Option A is missing."
                )


            if not option_b:

                raise ValueError(
                    "Option B is missing."
                )


            if not option_c:

                raise ValueError(
                    "Option C is missing."
                )


            if not option_d:

                raise ValueError(
                    "Option D is missing."
                )


            if not correct_option:

                raise ValueError(
                    "Correct option is missing."
                )


            cursor.execute(
                """
                INSERT INTO Questions
                (
                    exam_id,
                    question,
                    option_a,
                    option_b,
                    option_c,
                    option_d,
                    correct_option
                )
                VALUES
                (
                    ?, ?, ?, ?, ?, ?, ?
                )
                """,
                (
                    exam_id,
                    question,
                    option_a,
                    option_b,
                    option_c,
                    option_d,
                    correct_option
                )
            )


        conn.commit()


        return jsonify({

            "status":
                "success",

            "message":
                "Quiz saved successfully.",

            "exam_id":
                exam_id,

            "total_questions":
                total_questions

        }), 201


    except Exception as e:

        if conn:

            conn.rollback()

        print(
            "Admin save quiz error:",
            e
        )

        return jsonify({

            "status":
                "error",

            "message":
                "Unable to save quiz.",

            "error":
                str(e)

        }), 500


    finally:

        if conn:

            conn.close()

@admin_dashboard_bp.route(
    "/api/live-frame/<int:candidate_id>/<int:exam_id>"
)
@admin_required
def live_frame(
    candidate_id,
    exam_id
):

    try:

        base_dir = Path(
            __file__
        ).resolve().parent.parent

        filepath = (
            base_dir
            / "evidence"
            / f"candidate_{candidate_id}"
            / f"exam_{exam_id}"
            / "live"
            / "latest.jpg"
        )

        if not filepath.exists():

            return jsonify({

                "success":
                    False,

                "message":
                    "No live frame available."

            }), 404

        return send_file(
            filepath,
            mimetype="image/jpeg"
        )

    except Exception as error:

        return jsonify({

            "success":
                False,

            "message":
                str(error)

        }), 500


# ---------------------------------------------------------
# VIOLATIONS & EVIDENCE PAGE
# ---------------------------------------------------------

# ---------------------------------------------------------
# VIOLATIONS & EVIDENCE PAGE
# ---------------------------------------------------------

@admin_dashboard_bp.route(
    "/violations-evidence",
    methods=["GET"]
)
@admin_login_required
def violations_evidence_page():
    return render_template(
        "violations_evidence.html"
    )


# ---------------------------------------------------------
# INTEGRITY ANALYSIS PAGE
# ---------------------------------------------------------

@admin_dashboard_bp.route(
    "/integrity-analysis",
    methods=["GET"]
)
@admin_login_required
def integrity_analysis_page():
    return render_template(
        "integrity_analysis.html"
    )


# ---------------------------------------------------------
# INTEGRITY ANALYTICS PAGE
# ---------------------------------------------------------

@admin_dashboard_bp.route(
    "/integrity-analytics",
    methods=["GET"]
)
@admin_login_required
def integrity_analytics_page():
    return render_template(
        "integrity_analytics.html"
    )


# ---------------------------------------------------------
# ENVIRONMENT CHECK PAGE
# ---------------------------------------------------------

@admin_dashboard_bp.route(
    "/environment-check",
    methods=["GET"]
)
@admin_login_required
def environment_check_page():
    return render_template(
        "environment_check.html"
    )


# ---------------------------------------------------------
# HELP SUPPORT PAGE
# ---------------------------------------------------------

@admin_dashboard_bp.route(
    "/help-support",
    methods=["GET"]
)
@admin_login_required
def help_support_page():
    return render_template(
        "help_support.html"
    )


# ---------------------------------------------------------
# SETTINGS PAGE
# ---------------------------------------------------------

@admin_dashboard_bp.route(
    "/settings",
    methods=["GET"]
)
@admin_login_required
def settings_page():
    return render_template(
        "settings.html"
    )


# ---------------------------------------------------------
# SECURE EVIDENCE IMAGE DELIVERY API
# ---------------------------------------------------------

@admin_dashboard_bp.route("/api/evidence", methods=["GET"])
@admin_dashboard_bp.route("/evidence/<path:filename>", methods=["GET"])
@admin_api_required
def serve_evidence(filename=None):
    raw_path = filename or request.args.get("path", "")
    if not raw_path:
        return jsonify({"success": False, "error": "Evidence path is required."}), 400

    # Normalize Windows backslashes and strip prefix
    normalized = str(raw_path).replace("\\", "/").strip()
    if normalized.startswith("/"):
        normalized = normalized.lstrip("/")
    if normalized.startswith("evidence/"):
        normalized = normalized[len("evidence/"):]

    evidence_root = (BASE_DIR / "evidence").resolve()
    target_file = (evidence_root / normalized).resolve()

    # Path traversal protection
    try:
        target_file.relative_to(evidence_root)
    except ValueError:
        return jsonify({"success": False, "error": "Invalid evidence path."}), 403

    if not target_file.exists() or not target_file.is_file():
        return jsonify({"success": False, "error": "Evidence file not found."}), 404

    suffix = target_file.suffix.lower()
    mimetypes = {
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".png": "image/png",
        ".webp": "image/webp"
    }
    mimetype = mimetypes.get(suffix, "image/jpeg")

    return send_file(target_file, mimetype=mimetype)


# ---------------------------------------------------------
# VIOLATIONS & EVIDENCE API
# ---------------------------------------------------------

@admin_dashboard_bp.route(
    "/api/violations-evidence",
    methods=["GET"]
)
@admin_api_required
def violations_evidence_api():

    conn = get_db_connection()

    try:

        # =================================================
        # FILTERS
        # =================================================

        search = request.args.get(
            "search",
            "",
            type=str
        ).strip()

        exam_id = request.args.get(
            "exam_id",
            type=int
        )

        category = request.args.get(
            "category",
            "",
            type=str
        ).strip()

        severity = request.args.get(
            "severity",
            "",
            type=str
        ).strip()

        page = request.args.get(
            "page",
            1,
            type=int
        )

        per_page = request.args.get(
            "per_page",
            10,
            type=int
        )

        if page < 1:
            page = 1

        if per_page < 1:
            per_page = 10

        offset = (page - 1) * per_page

        # =================================================
        # CATEGORY DEFINITIONS
        # =================================================

        categories = {
            "TAB_SWITCH": {
                "title": "Tab Switches",
                "description": "Browser activity"
            },

            "FOCUS_LOSS": {
                "title": "Focus Loss",
                "description": "Window focus events"
            },

            "FACE_ABSENCE": {
                "title": "Face Absence",
                "description": "Face presence events"
            },

            "FULLSCREEN_EXIT": {
                "title": "Fullscreen Exit",
                "description": "Fullscreen violations"
            },

            "COPY_PASTE": {
                "title": "Copy / Paste",
                "description": "Clipboard activity"
            },

            "SCREENSHOT": {
                "title": "Screenshots",
                "description": "Screenshot attempts"
            },

            "RIGHT_CLICK": {
                "title": "Right Click",
                "description": "Context menu activity"
            },

            "IDENTITY_MISMATCH": {
                "title": "Identity Mismatch",
                "description": "Verification failures"
            },

            "MULTIPLE_FACES": {
                "title": "Multiple Faces",
                "description": "Multiple person detection"
            },

            "OTHER": {
                "title": "Other Violations",
                "description": "Other suspicious events"
            }
        }

        # =================================================
        # DATABASE COLUMN CHECK
        # =================================================

        violation_columns = {
            row["name"]
            for row in conn.execute(
                "PRAGMA table_info(ViolationLogs)"
            ).fetchall()
        }

        has_severity = "severity" in violation_columns
        has_status = "status" in violation_columns

        # =================================================
        # CATEGORY SQL CONDITIONS
        # =================================================

        category_conditions = {
            "TAB_SWITCH": [
                "TAB_SWITCH",
                "TAB_SWITCHED",
                "TAB_SWITCH_DETECTED"
            ],

            "FOCUS_LOSS": [
                "FOCUS_LOSS",
                "FOCUS_LOST"
            ],

            "FACE_ABSENCE": [
                "FACE_ABSENCE",
                "NO_FACE",
                "FACE_NOT_DETECTED",
                "FACE_ABSENT"
            ],

            "FULLSCREEN_EXIT": [
                "FULLSCREEN_EXIT",
                "FULLSCREEN_EXITED"
            ],

            "COPY_PASTE": [
                "COPY",
                "PASTE",
                "COPY_PASTE"
            ],

            "SCREENSHOT": [
                "SCREENSHOT",
                "SCREENSHOT_ATTEMPT"
            ],

            "RIGHT_CLICK": [
                "RIGHT_CLICK",
                "CONTEXT_MENU"
            ],

            "IDENTITY_MISMATCH": [
                "IDENTITY_MISMATCH",
                "IDENTITY_FAILED",
                "UNKNOWN_FACE"
            ],

            "MULTIPLE_FACES": [
                "MULTIPLE_FACES",
                "MULTIPLE_FACE"
            ]
        }

        # =================================================
        # HELPER
        # =================================================

        def category_for_type(value):

            value = (
                str(value or "")
                .strip()
                .upper()
            )

            for key, values in category_conditions.items():

                for item in values:

                    if item in value:
                        return key

            return "OTHER"

        # =================================================
        # HELPER - SEVERITY
        # =================================================

        def calculate_severity(value):

            value = (
                str(value or "")
                .upper()
            )

            if (
                "MULTIPLE" in value
                or "IDENTITY" in value
                or "UNKNOWN" in value
                or "NO_FACE" in value
            ):
                return "High"

            if (
                "TAB" in value
                or "SCREENSHOT" in value
                or "FULLSCREEN" in value
            ):
                return "Medium"

            return "Low"

        # =================================================
        # GET ALL VIOLATIONS
        # =================================================

        base_query = """
            SELECT
                v.id,
                v.candidate_id,
                v.exam_id,
                v.violation_type,
                v.evidence_image,
                v.face_count,
                v.violation_time,

                c.name AS candidate_name,
                c.email AS candidate_email,

                e.title AS exam_title

            FROM ViolationLogs v

            LEFT JOIN Candidates c
                ON c.id = v.candidate_id

            LEFT JOIN Exams e
                ON e.id = v.exam_id

            WHERE 1 = 1
        """

        params = []

        # =================================================
        # SEARCH
        # =================================================

        if search:

            base_query += """
                AND (
                    c.name LIKE ?
                    OR c.email LIKE ?
                    OR CAST(v.candidate_id AS TEXT) LIKE ?
                    OR e.title LIKE ?
                )
            """

            search_value = f"%{search}%"

            params.extend([
                search_value,
                search_value,
                search_value,
                search_value
            ])

        # =================================================
        # EXAM FILTER
        # =================================================

        if exam_id is not None:

            base_query += """
                AND v.exam_id = ?
            """

            params.append(exam_id)

        # =================================================
        # GET RECORDS
        # =================================================

        base_query += """
            ORDER BY
                datetime(v.violation_time) DESC
        """

        rows = conn.execute(
            base_query,
            params
        ).fetchall()

        # =================================================
        # CONVERT DATABASE ROWS
        # =================================================

        all_events = []

        for row in rows:

            violation_type = row["violation_type"]

            event_category = category_for_type(
                violation_type
            )

            if (
                category
                and category != "ALL"
                and event_category != category
            ):
                continue

            if has_severity:

                db_severity = conn.execute(
                    """
                    SELECT severity
                    FROM ViolationLogs
                    WHERE id = ?
                    """,
                    (row["id"],)
                ).fetchone()

                event_severity = (
                    db_severity["severity"]
                    if db_severity
                    else calculate_severity(
                        violation_type
                    )
                )

            else:

                event_severity = calculate_severity(
                    violation_type
                )

            if (
                severity
                and severity != "ALL"
                and event_severity.upper()
                != severity.upper()
            ):
                continue

            if has_status:

                status_row = conn.execute(
                    """
                    SELECT status
                    FROM ViolationLogs
                    WHERE id = ?
                    """,
                    (row["id"],)
                ).fetchone()

                event_status = (
                    status_row["status"]
                    if status_row
                    else "Detected"
                )

            else:

                event_status = "Detected"

            all_events.append({

                "id": row["id"],

                "candidate_id":
                    row["candidate_id"],

                "candidate_name":
                    row["candidate_name"] or "Unknown",

                "candidate_email":
                    row["candidate_email"] or "",

                "exam_id":
                    row["exam_id"],

                "exam_title":
                    row["exam_title"] or "Unknown Examination",

                "violation_type":
                    violation_type or "Other",

                "category":
                    event_category,

                "severity":
                    event_severity,

                "detected_at":
                    row["violation_time"],

                "evidence":
                    row["evidence_image"],

                "face_count":
                    row["face_count"] or 0,

                "status":
                    event_status

            })

        # =================================================
        # PAGINATION
        # =================================================

        total_events = len(all_events)

        total_pages = max(
            1,
            (
                total_events
                + per_page
                - 1
            ) // per_page
        )

        paginated_events = all_events[
            offset:
            offset + per_page
        ]

        # =================================================
        # CATEGORY COUNTS
        # =================================================

        category_counts = {}

        for key in categories:

            category_counts[key] = sum(
                1
                for event in all_events
                if event["category"] == key
            )

        # =================================================
        # SEVERITY COUNTS
        # =================================================

        high_count = sum(
            1
            for event in all_events
            if event["severity"] == "High"
        )

        medium_count = sum(
            1
            for event in all_events
            if event["severity"] == "Medium"
        )

        low_count = sum(
            1
            for event in all_events
            if event["severity"] == "Low"
        )

        # =================================================
        # EVIDENCE
        # =================================================

        evidence_events = [
            event
            for event in all_events
            if event["evidence"]
        ]

        total_evidence = len(
            evidence_events
        )

        # =================================================
        # RECENT EVIDENCE
        # =================================================

        recent_evidence = evidence_events[:6]

        # =================================================
        # TIMELINE
        # =================================================

        timeline = all_events[:10]

        # =================================================
        # EXAMINATION LIST
        # =================================================

        exams = conn.execute(
            """
            SELECT
                id,
                title
            FROM Exams
            ORDER BY
                datetime(start_time) DESC
            """
        ).fetchall()

        exam_list = [
            {
                "id": exam["id"],
                "title": exam["title"]
            }
            for exam in exams
        ]

        # =================================================
        # RESPONSE
        # =================================================

        return jsonify({

            "success": True,

            "statistics": {

                "tab_switches":
                    category_counts["TAB_SWITCH"],

                "focus_loss":
                    category_counts["FOCUS_LOSS"],

                "face_absence":
                    category_counts["FACE_ABSENCE"],

                "fullscreen_exit":
                    category_counts["FULLSCREEN_EXIT"],

                "copy_paste":
                    category_counts["COPY_PASTE"],

                "screenshots":
                    category_counts["SCREENSHOT"],

                "right_click":
                    category_counts["RIGHT_CLICK"],

                "identity_mismatch":
                    category_counts["IDENTITY_MISMATCH"],

                "multiple_faces":
                    category_counts["MULTIPLE_FACES"],

                "other":
                    category_counts["OTHER"],

                "total":
                    total_events,

                "high":
                    high_count,

                "medium":
                    medium_count,

                "low":
                    low_count,

                "total_evidence":
                    total_evidence

            },

            "events":
                paginated_events,

            "recent_evidence":
                recent_evidence,

            "timeline":
                timeline,

            "exams":
                exam_list,

            "pagination": {

                "page":
                    page,

                "per_page":
                    per_page,

                "total":
                    total_events,

                "total_pages":
                    total_pages

            }

        })

    except Exception as error:

        print(
            "Violations & Evidence error:",
            repr(error)
        )

        return jsonify({

            "success": False,

            "message":
                "Unable to load violations and evidence.",

            "error":
                str(error)

        }), 500

    finally:

        conn.close()


# ============================================================
# INTEGRITY ANALYSIS CENTER
# ============================================================

# ------------------------------------------------------------
# SCORING CONFIGURATION
# ------------------------------------------------------------

EVENT_WEIGHTS = {
    "TAB_SWITCH": 4,
    "TAB_SWITCHES": 4,

    "FOCUS_LOSS": 3,
    "FOCUS": 3,

    "FACE_ABSENCE": 8,
    "NO_FACE": 8,

    "FULLSCREEN_EXIT": 4,
    "FULLSCREEN": 4,

    "COPY_PASTE": 5,
    "COPY": 5,
    "PASTE": 5,

    "SCREENSHOT": 6,

    "RIGHT_CLICK": 2,

    "IDENTITY_MISMATCH": 10,
    "UNKNOWN_FACE": 10,

    "MULTIPLE_FACES": 10,
    "MULTIPLE_FACE": 10,

    "OTHER": 3
}


SEVERITY_WEIGHTS = {
    "LOW": 1,
    "MEDIUM": 3,
    "HIGH": 5
}


def normalize_violation_type(value):

    if not value:
        return "OTHER"

    value = str(value).upper().strip()

    value = value.replace("-", "_")
    value = value.replace(" ", "_")

    return value


def get_event_weight(violation_type):

    violation_type = normalize_violation_type(
        violation_type
    )

    for key, weight in EVENT_WEIGHTS.items():

        if key in violation_type:
            return weight

    return EVENT_WEIGHTS["OTHER"]


def get_severity(violation_type):

    violation_type = normalize_violation_type(
        violation_type
    )

    if any(
        key in violation_type
        for key in [
            "IDENTITY",
            "UNKNOWN",
            "MULTIPLE",
            "NO_FACE",
            "FACE_ABSENCE"
        ]
    ):
        return "HIGH"

    if any(
        key in violation_type
        for key in [
            "TAB",
            "FOCUS",
            "SCREENSHOT",
            "COPY",
            "PASTE",
            "FULLSCREEN"
        ]
    ):
        return "MEDIUM"

    return "LOW"


def calculate_integrity_score(
    events,
    face_presence_ratio=1.0,
    duration_minutes=1
):

    total_events = len(events)

    if total_events == 0:

        score = (
            100 *
            max(
                0,
                min(1, face_presence_ratio)
            )
        )

        return {
            "score": round(score, 2),
            "event_penalty": 0,
            "severity_penalty": 0,
            "face_penalty": round(
                (1 - face_presence_ratio) * 30,
                2
            ),
            "risk": (
                "Low"
                if score >= 80
                else "Medium"
                if score >= 60
                else "High"
            )
        }

    # --------------------------------------------------------
    # EVENT FREQUENCY
    # --------------------------------------------------------

    duration_minutes = max(
        float(duration_minutes),
        1
    )

    events_per_minute = (
        total_events /
        duration_minutes
    )

    # Frequency penalty is capped.
    event_penalty = min(
        30,
        events_per_minute * 8
    )

    # --------------------------------------------------------
    # SEVERITY PENALTY
    # --------------------------------------------------------

    severity_points = 0

    for event in events:

        severity = get_severity(
            event.get("violation_type")
        )

        event_weight = get_event_weight(
            event.get("violation_type")
        )

        severity_points += (
            event_weight *
            SEVERITY_WEIGHTS[severity]
        )

    severity_penalty = min(
        40,
        severity_points * 0.7
    )

    # --------------------------------------------------------
    # FACE PRESENCE PENALTY
    # --------------------------------------------------------

    face_presence_ratio = max(
        0,
        min(
            1,
            float(face_presence_ratio)
        )
    )

    face_penalty = (
        1 -
        face_presence_ratio
    ) * 30

    # --------------------------------------------------------
    # FINAL SCORE
    # --------------------------------------------------------

    raw_score = (
        100
        - event_penalty
        - severity_penalty
        - face_penalty
    )

    score = max(
        0,
        min(
            100,
            raw_score
        )
    )

    if score >= 80:
        risk = "Low"

    elif score >= 60:
        risk = "Medium"

    else:
        risk = "High"

    return {
        "score": round(score, 2),
        "event_penalty": round(
            event_penalty,
            2
        ),
        "severity_penalty": round(
            severity_penalty,
            2
        ),
        "face_penalty": round(
            face_penalty,
            2
        ),
        "risk": risk
    }


def calculate_face_presence(events):

    face_events = [
        e
        for e in events
        if e.get("face_count") is not None
    ]

    if not face_events:
        return 1.0

    present = sum(
        1
        for e in face_events
        if int(e.get("face_count") or 0) == 1
    )

    return present / len(face_events)


# ============================================================
# API - INTEGRITY ANALYSIS
# ============================================================

# ============================================================
# INTEGRITY ANALYSIS API
# ============================================================

@admin_dashboard_bp.route(
    "/api/integrity-analysis",
    methods=["GET"]
)
@admin_required
def integrity_analysis_api():

    conn = None

    try:

        # ============================================================
        # IMPORTS
        # ============================================================

        # ============================================================
        # DATABASE CONNECTION
        # ============================================================

        conn = get_db_connection()
        conn.row_factory = sqlite3.Row

        cursor = conn.cursor()

        # ============================================================
        # OPTIONAL EXAM FILTER
        # ============================================================

        exam_id = request.args.get(
            "exam_id",
            ""
        ).strip()

        # ============================================================
        # PHOTO DIRECTORY
        #
        # Project:
        # C:\Projects\online-exam-monitoring-integrity-analytics-platform
        #
        # Photos:
        # static/uploads/photos/
        # ============================================================

        PHOTO_DIR = os.path.join(
            current_app.root_path,
            "static",
            "uploads",
            "photos"
        )

        # ============================================================
        # HELPER:
        # FIND CANDIDATE PHOTO
        #
        # Example email:
        # prashanthikoda03@gmail.com
        #
        # Example file:
        # prashanthikoda03_at_gmail_com_bc418bcc.png
        #
        # We search using the email-derived prefix so the UUID/hash
        # at the end of the filename does not matter.
        # ============================================================

        def get_candidate_photo(email):

            if not email:
                return "/static/images/default-avatar.png"

            email = str(email).strip().lower()

            # Convert:
            # prashanthikoda03@gmail.com
            #
            # to:
            # prashanthikoda03_at_gmail_com

            email_prefix = (
                email
                .replace("@", "_at_")
                .replace(".", "_")
                .replace(" ", "_")
            )

            # Search all supported image types
            patterns = [
                os.path.join(
                    PHOTO_DIR,
                    email_prefix + "_*.png"
                ),
                os.path.join(
                    PHOTO_DIR,
                    email_prefix + "_*.jpg"
                ),
                os.path.join(
                    PHOTO_DIR,
                    email_prefix + "_*.jpeg"
                ),
                os.path.join(
                    PHOTO_DIR,
                    email_prefix + "_*.webp"
                ),

                # Also support exact filename without UUID
                os.path.join(
                    PHOTO_DIR,
                    email_prefix + ".png"
                ),
                os.path.join(
                    PHOTO_DIR,
                    email_prefix + ".jpg"
                ),
                os.path.join(
                    PHOTO_DIR,
                    email_prefix + ".jpeg"
                ),
                os.path.join(
                    PHOTO_DIR,
                    email_prefix + ".webp"
                )
            ]

            for pattern in patterns:

                matches = glob.glob(pattern)

                if matches:

                    # Take the first matching image
                    photo_file = matches[0]

                    filename = os.path.basename(
                        photo_file
                    )

                    # Browser URL
                    return (
                        "/static/uploads/photos/"
                        + filename
                    )

            # ========================================================
            # FALLBACK
            # ========================================================

            return "/static/images/default-avatar.png"

        # ============================================================
        # EXAMINATION LIST
        # ============================================================

        cursor.execute("""
            SELECT
                id,
                title,
                topic,
                difficulty
            FROM Exams
            ORDER BY id DESC
        """)

        exams = [
            dict(row)
            for row in cursor.fetchall()
        ]

        # ============================================================
        # INTEGRITY SCORE DATA
        # ============================================================

        query = """
            SELECT
                i.id,
                i.candidate_id,
                i.exam_id,

                c.name AS candidate_name,
                c.email AS candidate_email,

                e.title AS exam_title,
                e.topic AS exam_topic,

                i.integrity_score,
                i.face_presence_ratio,
                i.warning_count,
                i.risk_label,
                i.generated_at

            FROM IntegrityScores i

            INNER JOIN Candidates c
                ON c.id = i.candidate_id

            INNER JOIN Exams e
                ON e.id = i.exam_id
        """

        params = []

        # ============================================================
        # EXAM FILTER
        # ============================================================

        if exam_id:

            try:

                exam_id_value = int(
                    exam_id
                )

            except ValueError:

                return jsonify({
                    "status": "error",
                    "message": "Invalid examination ID."
                }), 400

            query += """
                WHERE i.exam_id = ?
            """

            params.append(
                exam_id_value
            )

        # ============================================================
        # ORDER
        # ============================================================

        query += """
            ORDER BY i.generated_at DESC
        """

        cursor.execute(
            query,
            params
        )

        integrity_rows = cursor.fetchall()

        # ============================================================
        # SESSION DATA
        # ============================================================

        sessions = []

        for row in integrity_rows:

            candidate_id = row[
                "candidate_id"
            ]

            current_exam_id = row[
                "exam_id"
            ]

            candidate_email = (
                row["candidate_email"]
                or ""
            )

            # ========================================================
            # CANDIDATE PHOTO
            # ========================================================

            candidate_photo = get_candidate_photo(
                candidate_email
            )

            # ========================================================
            # COUNT VIOLATION EVENTS
            # ========================================================

            cursor.execute("""
                SELECT
                    COUNT(*) AS total_events

                FROM ViolationLogs

                WHERE candidate_id = ?
                AND exam_id = ?
            """, (
                candidate_id,
                current_exam_id
            ))

            event_row = cursor.fetchone()

            total_events = (
                event_row["total_events"]
                if event_row
                else 0
            )

            # ========================================================
            # VIOLATION SEVERITY
            # ========================================================

            cursor.execute("""
                SELECT
                    violation_type,
                    COUNT(*) AS count

                FROM ViolationLogs

                WHERE candidate_id = ?
                AND exam_id = ?

                GROUP BY violation_type
            """, (
                candidate_id,
                current_exam_id
            ))

            violation_rows = cursor.fetchall()

            severity_score = 0

            severity_weights = {

                "MULTIPLE_FACES": 20,

                "Multiple faces detected.": 20,

                "NO_FACE": 10,

                "No face detected.": 10,

                "UNKNOWN_FACE": 15,

                "Identity mismatch.": 15,

                "Window lost focus": 8,

                "TAB_SWITCH": 8,

                "Tab switched": 8,

                "FACE_MISMATCH": 15,

                "Face mismatch": 15
            }

            for violation in violation_rows:

                violation_type = (
                    violation["violation_type"]
                    or ""
                ).strip()

                count = (
                    violation["count"]
                    or 0
                )

                # Case-insensitive lookup
                normalized_type = (
                    violation_type.upper()
                )

                weight = 5

                if normalized_type in (
                    "MULTIPLE_FACES",
                    "MULTIPLE FACES DETECTED."
                ):
                    weight = 20

                elif normalized_type in (
                    "NO_FACE",
                    "NO FACE DETECTED."
                ):
                    weight = 10

                elif normalized_type in (
                    "UNKNOWN_FACE",
                    "IDENTITY MISMATCH.",
                    "FACE_MISMATCH",
                    "FACE MISMATCH"
                ):
                    weight = 15

                elif normalized_type in (
                    "WINDOW LOST FOCUS",
                    "TAB_SWITCH",
                    "TAB SWITCHED"
                ):
                    weight = 8

                severity_score += (
                    weight * count
                )

            # ========================================================
            # FACE PRESENCE
            # ========================================================

            try:

                face_presence_ratio = float(
                    row["face_presence_ratio"]
                    or 0
                )

            except (
                TypeError,
                ValueError
            ):

                face_presence_ratio = 0

            # ========================================================
            # HANDLE RATIO
            #
            # If database stores:
            # 0.95 -> 95%
            #
            # If database stores:
            # 95 -> 95%
            # ========================================================

            if (
                0 <= face_presence_ratio <= 1
            ):

                face_presence = (
                    face_presence_ratio * 100
                )

            else:

                face_presence = (
                    face_presence_ratio
                )

            face_presence = min(
                max(
                    face_presence,
                    0
                ),
                100
            )

            # ========================================================
            # INTEGRITY SCORE
            # ========================================================

            try:

                integrity_score = float(
                    row["integrity_score"]
                    or 0
                )

            except (
                TypeError,
                ValueError
            ):

                integrity_score = 0

            integrity_score = min(
                max(
                    integrity_score,
                    0
                ),
                100
            )

            # ========================================================
            # RISK LABEL
            # ========================================================

            risk_label = (
                row["risk_label"]
                or ""
            ).strip()

            if not risk_label:

                if integrity_score >= 80:
                    risk_label = "LOW"

                elif integrity_score >= 60:
                    risk_label = "MEDIUM"

                else:
                    risk_label = "HIGH"

            # ========================================================
            # SESSION OBJECT
            # ========================================================

            sessions.append({

                "id":
                    row["id"],

                "candidateId":
                    candidate_id,

                "candidateName":
                    row["candidate_name"]
                    or "Unknown Candidate",

                "candidateEmail":
                    candidate_email,

                # IMPORTANT
                # This is the path your frontend needs.
                "photo":
                    candidate_photo,

                "examId":
                    current_exam_id,

                "examTitle":
                    row["exam_title"]
                    or "Unknown Examination",

                "examTopic":
                    row["exam_topic"]
                    or "",

                "totalEvents":
                    total_events,

                "severityScore":
                    round(
                        severity_score,
                        1
                    ),

                "facePresence":
                    round(
                        face_presence,
                        1
                    ),

                "integrityScore":
                    round(
                        integrity_score,
                        1
                    ),

                "riskLevel":
                    risk_label,

                "warningCount":
                    row["warning_count"]
                    or 0,

                "generatedAt":
                    row["generated_at"]
            })

        # ============================================================
        # SUMMARY
        # ============================================================

        total_sessions = len(
            sessions
        )

        if total_sessions > 0:

            average_score = (
                sum(
                    s["integrityScore"]
                    for s in sessions
                )
                / total_sessions
            )

            average_face = (
                sum(
                    s["facePresence"]
                    for s in sessions
                )
                / total_sessions
            )

        else:

            average_score = 0
            average_face = 0

        # ============================================================
        # RISK COUNTS
        # ============================================================

        low_risk = sum(
            1
            for s in sessions
            if str(
                s["riskLevel"]
            ).upper() == "LOW"
        )

        medium_risk = sum(
            1
            for s in sessions
            if str(
                s["riskLevel"]
            ).upper() == "MEDIUM"
        )

        high_risk = sum(
            1
            for s in sessions
            if str(
                s["riskLevel"]
            ).upper() == "HIGH"
        )

        # ============================================================
        # SUMMARY OBJECT
        # ============================================================

        summary = {

            "totalSessions":
                total_sessions,

            "averageScore":
                round(
                    average_score,
                    1
                ),

            "lowRisk":
                low_risk,

            "mediumRisk":
                medium_risk,

            "highRisk":
                high_risk,

            "facePresence":
                round(
                    average_face,
                    1
                )
        }

        # ============================================================
        # SCORE DISTRIBUTION
        # ============================================================

        distribution = {

            "labels": [
                "0-20",
                "21-40",
                "41-60",
                "61-80",
                "81-100"
            ],

            "values": [
                0,
                0,
                0,
                0,
                0
            ]
        }

        for session in sessions:

            score = session[
                "integrityScore"
            ]

            if score <= 20:

                distribution[
                    "values"
                ][0] += 1

            elif score <= 40:

                distribution[
                    "values"
                ][1] += 1

            elif score <= 60:

                distribution[
                    "values"
                ][2] += 1

            elif score <= 80:

                distribution[
                    "values"
                ][3] += 1

            else:

                distribution[
                    "values"
                ][4] += 1

        # ============================================================
        # RISK DISTRIBUTION
        # ============================================================

        risk_distribution = {

            "labels": [
                "Low",
                "Medium",
                "High"
            ],

            "values": [
                low_risk,
                medium_risk,
                high_risk
            ]
        }

        # ============================================================
        # HEATMAP
        #
        # IMPORTANT:
        #
        # Your JS expects:
        #
        # heatmap.events
        #
        # Each event has 24 hourly values.
        #
        # Example:
        #
        # {
        #   "event": "NO_FACE",
        #   "values": [0,0,1,0,...]
        # }
        # ============================================================

        heatmap_query = """
            SELECT
                violation_type,
                violation_time

            FROM ViolationLogs
        """

        heatmap_params = []

        if exam_id:

            heatmap_query += """
                WHERE exam_id = ?
            """

            heatmap_params.append(
                int(exam_id)
            )

        heatmap_query += """
            ORDER BY violation_time ASC
        """

        cursor.execute(
            heatmap_query,
            heatmap_params
        )

        heatmap_rows = cursor.fetchall()

        # ============================================================
        # BUILD HEATMAP BY VIOLATION TYPE
        # ============================================================

        heatmap_events = {}

        for row in heatmap_rows:

            event_name = (
                row["violation_type"]
                or "OTHER"
            ).strip()

            if event_name not in heatmap_events:

                heatmap_events[
                    event_name
                ] = [0] * 24

            violation_time = (
                row["violation_time"]
            )

            if not violation_time:
                continue

            # --------------------------------------------------------
            # Try to extract hour from timestamp.
            # --------------------------------------------------------

            try:

                from datetime import datetime

                parsed_time = None

                time_string = str(
                    violation_time
                ).strip()

                datetime_formats = [

                    "%Y-%m-%d %H:%M:%S",

                    "%Y-%m-%d %H:%M",

                    "%Y/%m/%d %H:%M:%S",

                    "%Y/%m/%d %H:%M",

                    "%d-%m-%Y %H:%M:%S",

                    "%d-%m-%Y %H:%M"
                ]

                for fmt in datetime_formats:

                    try:

                        parsed_time = (
                            datetime.strptime(
                                time_string,
                                fmt
                            )
                        )

                        break

                    except ValueError:

                        continue

                # ----------------------------------------------------
                # If normal datetime parsing fails, try ISO format.
                # ----------------------------------------------------

                if parsed_time is None:

                    try:

                        parsed_time = (
                            datetime.fromisoformat(
                                time_string
                            )
                        )

                    except ValueError:

                        parsed_time = None

                if parsed_time:

                    hour = parsed_time.hour

                    heatmap_events[
                        event_name
                    ][hour] += 1

            except Exception:

                continue

        # ============================================================
        # CONVERT HEATMAP TO FRONTEND FORMAT
        # ============================================================

        heatmap_event_list = []

        for event_name, values in (
            heatmap_events.items()
        ):

            heatmap_event_list.append({

                "event":
                    event_name,

                "values":
                    values
            })

        heatmap = {

            "events":
                heatmap_event_list
        }

        # ============================================================
        # COHORT ANALYSIS
        # ============================================================

        cohort_query = """
            SELECT

                e.id AS exam_id,

                e.title AS exam_title,

                COUNT(i.id) AS sessions,

                AVG(i.integrity_score)
                    AS average_score,

                SUM(
                    CASE
                        WHEN UPPER(
                            i.risk_label
                        ) = 'HIGH'
                        THEN 1
                        ELSE 0
                    END
                ) AS high_risk

            FROM Exams e

            LEFT JOIN IntegrityScores i
                ON i.exam_id = e.id
        """

        cohort_params = []

        if exam_id:

            cohort_query += """
                WHERE e.id = ?
            """

            cohort_params.append(
                int(exam_id)
            )

        cohort_query += """
            GROUP BY
                e.id,
                e.title

            ORDER BY
                e.id DESC
        """

        cursor.execute(
            cohort_query,
            cohort_params
        )

        cohort_rows = cursor.fetchall()

        cohorts = []

        for row in cohort_rows:

            cohorts.append({

                "examId":
                    row["exam_id"],

                "cohort":
                    row["exam_title"]
                    or "Unknown Examination",

                "sessions":
                    row["sessions"]
                    or 0,

                "averageScore":
                    round(
                        float(
                            row["average_score"]
                            or 0
                        ),
                        1
                    ),

                "highRisk":
                    row["high_risk"]
                    or 0
            })

        # ============================================================
        # RECENT ACTIVITY
        # ============================================================

        activity_query = """
            SELECT

                v.id,

                v.candidate_id,

                c.name AS candidate_name,

                c.email AS candidate_email,

                v.exam_id,

                e.title AS exam_title,

                v.violation_type,

                v.face_count,

                v.violation_time

            FROM ViolationLogs v

            LEFT JOIN Candidates c
                ON c.id = v.candidate_id

            LEFT JOIN Exams e
                ON e.id = v.exam_id
        """

        activity_params = []

        if exam_id:

            activity_query += """
                WHERE v.exam_id = ?
            """

            activity_params.append(
                int(exam_id)
            )

        activity_query += """
            ORDER BY
                v.violation_time DESC

            LIMIT 10
        """

        cursor.execute(
            activity_query,
            activity_params
        )

        activity_rows = cursor.fetchall()

        recent_activity = []

        for row in activity_rows:

            recent_activity.append({

                "id":
                    row["id"],

                "candidateId":
                    row["candidate_id"],

                "candidateName":
                    row["candidate_name"]
                    or "Unknown Candidate",

                "candidateEmail":
                    row["candidate_email"]
                    or "",

                # Photo also available here
                "photo":
                    get_candidate_photo(
                        row["candidate_email"]
                    ),

                "examId":
                    row["exam_id"],

                "examTitle":
                    row["exam_title"]
                    or "Unknown Examination",

                "event":
                    row["violation_type"]
                    or "Integrity Event",

                "faceCount":
                    row["face_count"]
                    or 0,

                "time":
                    row["violation_time"]
                    or ""
            })

        # ============================================================
        # BEHAVIORAL CLUSTERS
        # ============================================================

        clusters = []

        if len(sessions) >= 3:

            try:

                from modules.datasciencemodule import (
                    perform_behavioral_clustering
                )

                clusters = (
                    perform_behavioral_clustering(
                        sessions
                    )
                )

                if not isinstance(
                    clusters,
                    list
                ):

                    clusters = []

            except Exception as cluster_error:

                print(
                    "Clustering warning:",
                    repr(cluster_error)
                )

                clusters = []

        # ============================================================
        # DEBUG INFORMATION
        #
        # This will help you verify photos in terminal.
        # ============================================================

        print(
            "Integrity Analysis:",
            len(sessions),
            "sessions loaded."
        )

        for session in sessions:

            print(
                "Candidate:",
                session["candidateName"],
                "| Email:",
                session["candidateEmail"],
                "| Photo:",
                session["photo"]
            )

        # ============================================================
        # FINAL RESPONSE
        # ============================================================

        return jsonify({

            "status":
                "success",

            "summary":
                summary,

            "sessions":
                sessions,

            "exams":
                exams,

            "distribution":
                distribution,

            "riskDistribution":
                risk_distribution,

            "heatmap":
                heatmap,

            "clusters":
                clusters,

            "cohorts":
                cohorts,

            "recentActivity":
                recent_activity

        }), 200

    # ================================================================
    # ERROR HANDLING
    # ================================================================

    except Exception as error:

        import traceback

        print(
            "Integrity analysis API error:",
            repr(error)
        )

        traceback.print_exc()

        return jsonify({

            "status":
                "error",

            "message":
                "Unable to load integrity analysis.",

            "error":
                str(error)

        }), 500

    # ================================================================
    # CLOSE DATABASE
    # ================================================================

    finally:

        if conn:

            conn.close()

@admin_dashboard_bp.route(
    "/api/integrity-analysis/export",
    methods=["GET"]
)
@admin_required
def export_integrity_analysis():

    try:

        # ============================================================
        # REUSE THE EXISTING INTEGRITY ANALYSIS API
        # ============================================================

        response = integrity_analysis_api()


        # Flask route functions can return:
        #
        # Response
        #
        # or:
        #
        # (Response, status_code)

        if isinstance(response, tuple):

            response_object = response[0]

            status_code = response[1]

        else:

            response_object = response

            status_code = 200


        if status_code != 200:

            return response


        data = response_object.get_json()


        if not data:

            return jsonify({

                "success": False,

                "message":
                    "No integrity analysis data available."

            }), 404


        # ============================================================
        # CSV GENERATION
        # ============================================================

        import csv
        import io

        output = io.StringIO()

        writer = csv.writer(
            output
        )


        # ============================================================
        # REPORT HEADER
        # ============================================================

        writer.writerow([
            "Integrity Analysis Report"
        ])

        writer.writerow([])


        # ============================================================
        # SUMMARY
        # ============================================================

        summary = (
            data.get(
                "summary",
                {}
            )
        )


        writer.writerow([
            "SUMMARY"
        ])

        writer.writerow([
            "Total Sessions",
            summary.get(
                "totalSessions",
                0
            )
        ])

        writer.writerow([
            "Average Integrity Score",
            summary.get(
                "averageScore",
                0
            )
        ])

        writer.writerow([
            "Low Risk",
            summary.get(
                "lowRisk",
                0
            )
        ])

        writer.writerow([
            "Medium Risk",
            summary.get(
                "mediumRisk",
                0
            )
        ])

        writer.writerow([
            "High Risk",
            summary.get(
                "highRisk",
                0
            )
        ])

        writer.writerow([
            "Average Face Presence",
            summary.get(
                "facePresence",
                0
            )
        ])

        writer.writerow([])


        # ============================================================
        # SESSION DATA
        # ============================================================

        writer.writerow([
            "SESSION DETAILS"
        ])


        writer.writerow([

            "Candidate",

            "Candidate ID",

            "Email",

            "Examination",

            "Total Events",

            "Severity Score",

            "Face Presence (%)",

            "Integrity Score",

            "Risk Level",

            "Warning Count",

            "Generated At"

        ])


        sessions = (
            data.get(
                "sessions",
                []
            )
        )


        for session_data in sessions:

            writer.writerow([

                session_data.get(
                    "candidateName",
                    ""
                ),

                session_data.get(
                    "candidateId",
                    ""
                ),

                session_data.get(
                    "candidateEmail",
                    ""
                ),

                session_data.get(
                    "examTitle",
                    ""
                ),

                session_data.get(
                    "totalEvents",
                    0
                ),

                session_data.get(
                    "severityScore",
                    0
                ),

                session_data.get(
                    "facePresence",
                    0
                ),

                session_data.get(
                    "integrityScore",
                    0
                ),

                session_data.get(
                    "riskLevel",
                    ""
                ),

                session_data.get(
                    "warningCount",
                    0
                ),

                session_data.get(
                    "generatedAt",
                    ""
                )

            ])


        # ============================================================
        # DOWNLOAD RESPONSE
        # ============================================================

        from flask import make_response

        csv_data = output.getvalue()

        response = make_response(
            csv_data
        )


        response.headers[
            "Content-Type"
        ] = (
            "text/csv; charset=utf-8"
        )


        response.headers[
            "Content-Disposition"
        ] = (
            "attachment; "
            "filename=integrity_analysis_report.csv"
        )


        return response


    except Exception as error:

        import traceback

        print(
            "Integrity export error:",
            repr(error)
        )

        traceback.print_exc()

        return jsonify({

            "success":
                False,

            "message":
                "Unable to export integrity analysis report.",

            "error":
                str(error)

        }), 500
# ==========================================================
# GET EXAMS (EXAMINATIONS / AI REPORTS)
# ==========================================================

@admin_dashboard_bp.route(
    "/api/ai-reports/exams",
    methods=["GET"]
)
@admin_api_required
def ai_reports_exams():

    conn = get_db_connection()

    try:

        rows = conn.execute("""
            SELECT
                id,
                title,
                topic,
                difficulty,
                duration,
                start_time,
                end_time
            FROM Exams
            ORDER BY id DESC
        """).fetchall()

        exams = []

        for row in rows:

            exams.append({
                "id": row["id"],
                "title": row["title"],
                "topic": row["topic"],
                "difficulty": row["difficulty"],
                "duration": row["duration"],
                "start_time": row["start_time"],
                "end_time": row["end_time"]
            })

        return jsonify({
            "success": True,
            "exams": exams
        })

    except Exception as error:

        return jsonify({
            "success": False,
            "error": str(error)
        }), 500

    finally:

        conn.close()


# ==========================================================
# DELETE EXAMINATION API
# ==========================================================

# ============================================================
# GET ALL EXAMINATIONS
# ============================================================
# ============================================================
# GET ALL EXAMINATIONS
# ============================================================

@admin_dashboard_bp.route(
    "/api/examinations",
    methods=["GET"]
)
@admin_api_required
def get_examinations():

    conn = get_db_connection()

    try:

        exams = conn.execute("""
            SELECT
                id,
                title,
                topic,
                difficulty,
                description,
                duration,
                total_questions,
                total_marks,
                start_time,
                end_time,
                created_at
            FROM Exams
            ORDER BY datetime(created_at) DESC, id DESC
        """).fetchall()

        examinations = []

        for exam in exams:

            examination = {
                "id": exam["id"],

                "title": (
                    exam["title"]
                    if exam["title"] is not None
                    else ""
                ),

                "topic": (
                    exam["topic"]
                    if exam["topic"] is not None
                    else ""
                ),

                "difficulty": (
                    exam["difficulty"]
                    if exam["difficulty"] is not None
                    else ""
                ),

                "description": (
                    exam["description"]
                    if exam["description"] is not None
                    else ""
                ),

                "duration": int(
                    exam["duration"] or 0
                ),

                "total_questions": int(
                    exam["total_questions"] or 0
                ),

                "total_marks": int(
                    exam["total_marks"] or 0
                ),

                "start_time": exam["start_time"],

                "end_time": exam["end_time"],

                "created_at": exam["created_at"]
            }

            examinations.append(
                examination
            )

        # Debug output
        print("\n===== EXAMS API RESPONSE =====")

        for examination in examinations:
            print(examination)

        print(
            "Total examinations:",
            len(examinations)
        )

        print(
            "===============================\n"
        )

        return jsonify({
            "success": True,
            "exams": examinations
        }), 200

    except Exception as error:

        print(
            "GET EXAMINATIONS ERROR:",
            error
        )

        return jsonify({
            "success": False,
            "message": "Failed to load examinations.",
            "error": str(error)
        }), 500

    finally:

        conn.close()


# ============================================================
# DELETE EXAMINATION
# ============================================================

@admin_dashboard_bp.route(
    "/api/examinations/<int:exam_id>",
    methods=["DELETE"]
)
@admin_api_required
def delete_examination(exam_id):

    conn = get_db_connection()

    try:

        exam = conn.execute(
            """
            SELECT
                id,
                title
            FROM Exams
            WHERE id = ?
            """,
            (exam_id,)
        ).fetchone()

        if not exam:

            return jsonify({
                "success": False,
                "message": "Examination not found."
            }), 404

        conn.execute(
            """
            DELETE FROM Exams
            WHERE id = ?
            """,
            (exam_id,)
        )

        conn.commit()

        print(
            f"Examination deleted: "
            f"{exam['id']} - {exam['title']}"
        )

        return jsonify({
            "success": True,
            "message": "Examination deleted successfully.",
            "exam_id": exam_id
        }), 200

    except Exception as error:

        conn.rollback()

        print(
            "DELETE EXAMINATION ERROR:",
            error
        )

        return jsonify({
            "success": False,
            "message": "Failed to delete examination.",
            "error": str(error)
        }), 500

    finally:

        conn.close()

# ============================================================
# DELETE EXAMINATION
# ============================================================

# ==========================================================
# GENERATE CANDIDATE AI REPORT
# ==========================================================

@admin_dashboard_bp.route(
    "/api/ai-reports/generate",
    methods=["POST"]
)
@admin_api_required
def generate_candidate_ai_report():

    data = request.get_json(
        silent=True
    )

    if not data:

        return jsonify({
            "success": False,
            "error": "Request body is required"
        }), 400

    candidate_id = data.get(
        "candidate_id"
    )

    exam_id = data.get(
        "exam_id"
    )

    if candidate_id is None or exam_id is None:

        return jsonify({
            "success": False,
            "error":
                "candidate_id and exam_id are required"
        }), 400

    conn = get_db_connection()

    try:

        # --------------------------------------------------
        # CHECK OLLAMA
        # --------------------------------------------------

        if not check_ollama():

            return jsonify({
                "success": False,
                "error":
                    "Ollama is not running. Start Ollama and try again."
            }), 503

        # --------------------------------------------------
        # CANDIDATE
        # --------------------------------------------------

        candidate = conn.execute("""
            SELECT
                id,
                name,
                email
            FROM Candidates
            WHERE id = ?
        """, (
            candidate_id,
        )).fetchone()

        if candidate is None:

            return jsonify({
                "success": False,
                "error": "Candidate not found"
            }), 404

        # --------------------------------------------------
        # EXAM
        # --------------------------------------------------

        exam = conn.execute("""
            SELECT
                id,
                title,
                topic,
                difficulty,
                description,
                duration,
                total_questions,
                total_marks
            FROM Exams
            WHERE id = ?
        """, (
            exam_id,
        )).fetchone()

        if exam is None:

            return jsonify({
                "success": False,
                "error": "Exam not found"
            }), 404

        # --------------------------------------------------
        # EXAM ATTEMPT
        # --------------------------------------------------

        attempt = conn.execute("""
            SELECT
                score,
                total_questions,
                percentage,
                result,
                submitted_at
            FROM ExamAttempts

            WHERE candidate_id = ?
              AND exam_id = ?

            LIMIT 1
        """, (
            candidate_id,
            exam_id
        )).fetchone()

        if attempt:

            attempt_data = dict(attempt)

        else:

            attempt_data = {

                "score": 0,

                "total_questions":
                    exam["total_questions"],

                "percentage": 0,

                "result":
                    "NOT SUBMITTED",

                "submitted_at": None
            }

        # --------------------------------------------------
        # INTEGRITY SCORE
        # --------------------------------------------------

        integrity = conn.execute("""
            SELECT
                integrity_score,
                face_presence_ratio,
                warning_count,
                risk_label,
                generated_at
            FROM IntegrityScores

            WHERE candidate_id = ?
              AND exam_id = ?

            LIMIT 1
        """, (
            candidate_id,
            exam_id
        )).fetchone()

        if integrity:

            integrity_data = dict(
                integrity
            )

        else:

            integrity_data = {

                "integrity_score": 0,

                "face_presence_ratio": 0,

                "warning_count": 0,

                "risk_label": "Unknown",

                "generated_at": None
            }

        # --------------------------------------------------
        # VIOLATIONS
        # --------------------------------------------------

        violation_rows = conn.execute("""
            SELECT
                violation_type,
                violation_time
            FROM ViolationLogs

            WHERE candidate_id = ?
              AND exam_id = ?

            ORDER BY violation_time ASC
        """, (
            candidate_id,
            exam_id
        )).fetchall()

        violations = {}

        for row in violation_rows:

            violation_type = (
                row["violation_type"]
                or "OTHER"
            ).strip()

            violations[violation_type] = (
                violations.get(
                    violation_type,
                    0
                ) + 1
            )

        # --------------------------------------------------
        # GENERATE AI REPORT
        # --------------------------------------------------

        report = generate_candidate_report(

            candidate=dict(candidate),

            exam=dict(exam),

            attempt=attempt_data,

            integrity=integrity_data,

            violations=violations

        )

        # --------------------------------------------------
        # SAVE HISTORY
        # --------------------------------------------------
        print(report)
        generated_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        cursor = conn.execute("""
            INSERT INTO AIReports (

                candidate_id,
                exam_id,
                report_type,
                title,
                risk_label,
                integrity_score,
                face_presence_ratio,
                warning_count,
                report_content,
                model_name,
                generated_at

            )

            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (

            candidate_id,

            exam_id,

            "candidate",

            f"Integrity Report - {candidate['name']}",

            integrity_data[
                "risk_label"
            ],

            integrity_data[
                "integrity_score"
            ],

            integrity_data[
                "face_presence_ratio"
            ],

            integrity_data[
                "warning_count"
            ],

            report,

            "ollama:" + (
                __import__(
                    "modules.ai_report_generator",
                    fromlist=[
                        "OLLAMA_MODEL"
                    ]
                ).OLLAMA_MODEL
            ),
            generated_at

        ))

        report_id = cursor.lastrowid

        conn.commit()

        return jsonify({

            "success": True,

            "report_id":
                report_id,

            "candidate": {
                "id":
                    candidate["id"],

                "name":
                    candidate["name"],

                "email":
                    candidate["email"]
            },

            "exam": {
                "id":
                    exam["id"],

                "title":
                    exam["title"]
            },

            "risk":
                integrity_data[
                    "risk_label"
                ],

            "integrity_score":
                integrity_data[
                    "integrity_score"
                ],

            "face_presence_ratio":
                integrity_data[
                    "face_presence_ratio"
                ],

            "warning_count":
                integrity_data[
                    "warning_count"
                ],

            "violations":
                violations,

            "assessment":
                report

        })

    except Exception as error:

        conn.rollback()

        print(
            "Candidate AI report error:",
            repr(error)
        )

        return jsonify({

            "success": False,

            "error":
                str(error)

        }), 500

    finally:

        conn.close()

@admin_dashboard_bp.route(
    "/api/ai-reports/generate-exam",
    methods=["POST"]
)
@admin_api_required
def generate_exam_ai_report():

    data = request.get_json()

    exam_id = data.get("exam_id")

    if not exam_id:

        return jsonify({
            "success": False,
            "error": "exam_id required"
        }), 400

    conn = get_db_connection()

    try:

        # ============================================================
        # 1. EXAM DETAILS
        # ============================================================

        exam = conn.execute("""

            SELECT *

            FROM Exams

            WHERE id=?

        """, (exam_id,)).fetchone()

        if not exam:

            return jsonify({
                "success": False,
                "error": "Exam not found"
            }), 404


        # ============================================================
        # 2. REGISTERED CANDIDATES
        # ============================================================

        registered = conn.execute("""

            SELECT COUNT(*) total

            FROM Candidates

        """).fetchone()["total"]


        # ============================================================
        # 3. ATTEMPT STATISTICS
        # ============================================================

        attempts = conn.execute("""

            SELECT

                COUNT(*) appeared,

                SUM(
                    CASE
                        WHEN result='PASS'
                        THEN 1
                        ELSE 0
                    END
                ) passed,

                SUM(
                    CASE
                        WHEN result='FAIL'
                        THEN 1
                        ELSE 0
                    END
                ) failed,

                AVG(score) average_score,

                AVG(percentage) average_percentage

            FROM ExamAttempts

            WHERE exam_id=?

        """, (exam_id,)).fetchone()


        # ============================================================
        # 4. INTEGRITY STATISTICS
        # ============================================================

        integrity = conn.execute("""

            SELECT

                AVG(integrity_score) average_integrity,

                AVG(face_presence_ratio) average_face,

                SUM(warning_count) total_warnings

            FROM IntegrityScores

            WHERE exam_id=?

        """, (exam_id,)).fetchone()


        # ============================================================
        # 5. RISK DISTRIBUTION
        # ============================================================

        risk_rows = conn.execute("""

            SELECT

                risk_label,

                COUNT(*) total

            FROM IntegrityScores

            WHERE exam_id=?

            GROUP BY risk_label

        """, (exam_id,)).fetchall()


        risk_distribution = {

            "Low": 0,

            "Medium": 0,

            "High": 0,

            "Unknown": 0

        }


        for row in risk_rows:

            risk_label = row["risk_label"] or "Unknown"

            if risk_label in risk_distribution:

                risk_distribution[risk_label] = row["total"]

            else:

                risk_distribution["Unknown"] += row["total"]


        # ============================================================
        # 6. VIOLATION SUMMARY
        # ============================================================

        rows = conn.execute("""

            SELECT

                violation_type,

                COUNT(*) total

            FROM ViolationLogs

            WHERE exam_id=?

            GROUP BY violation_type

        """, (exam_id,)).fetchall()


        violations = {}


        for row in rows:
            violation_type = row["violation_type"]
            normalized = violation_type.strip().upper()
            if normalized in (
                "MULTIPLE_FACES",
                "MULTIPLE FACES DETECTED."
            ):
                normalized = "Multiple Faces Detected"
            elif normalized in (
                "NO_FACE",
                "NO FACE DETECTED."
            ):
                normalized = "Face Missing"
            elif normalized == "WINDOW LOST FOCUS":
                normalized = "Window Focus Lost"
            violations[normalized] = (
                violations.get(normalized, 0) + row["total"]
            )
        # ============================================================
        # 7. CANDIDATE RANKING
        # ============================================================

        candidate_rows = conn.execute("""

            SELECT

                c.name,

                ea.score,

                ea.percentage,

                IFNULL(i.integrity_score,0) integrity_score,

                IFNULL(i.face_presence_ratio,0) face_presence,

                IFNULL(i.warning_count,0) warnings,

                IFNULL(i.risk_label,'Unknown') risk,

                COUNT(v.id) violations

            FROM ExamAttempts ea

            JOIN Candidates c

                ON ea.candidate_id=c.id

            LEFT JOIN IntegrityScores i

                ON ea.candidate_id=i.candidate_id

                AND ea.exam_id=i.exam_id

            LEFT JOIN ViolationLogs v

                ON ea.candidate_id=v.candidate_id

                AND ea.exam_id=v.exam_id

            WHERE ea.exam_id=?

            GROUP BY

                ea.candidate_id

            ORDER BY

                ea.percentage DESC,

                i.integrity_score DESC

        """, (exam_id,)).fetchall()


        candidates = []

        rank = 1


        for row in candidate_rows:

            candidates.append({

                "rank": rank,

                "name": row["name"],

                "score": row["score"],

                "percentage": row["percentage"],

                "integrity_score": row["integrity_score"],

                "face_presence": row["face_presence"],

                "warnings": row["warnings"],

                "risk": row["risk"],

                "violations": row["violations"]

            })

            rank += 1


        # ============================================================
        # 8. STATISTICS OBJECT
        # ============================================================

        statistics = {

            "registered_candidates": registered,

            "candidates": attempts["appeared"] or 0,

            "passed": attempts["passed"] or 0,

            "failed": attempts["failed"] or 0,

            "average_percentage": round(
                attempts["average_percentage"] or 0,
                2
            ),

            "average_integrity_score": round(
               integrity["average_integrity"] or 0,
               2
            ),

            "average_integrity_score": round(
                integrity["average_integrity"] or 0,
                2
            ),

            "average_face_presence": round(
                integrity["average_face"] or 0,
                2
            ),

            "total_violations": sum(
                violations.values()
            ),

            "evidence_count": sum(
                violations.values()
            )

        }


        # ============================================================
        # 9. GENERATE AI EXAM REPORT
        # ============================================================

        report = generate_exam_report(

            dict(exam),

            statistics,

            violations,

            risk_distribution,

            candidates

        )


        # ============================================================
        # 10. SAVE AI REPORT
        # ============================================================
        generated_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        cursor = conn.execute("""

            INSERT INTO AIReports(

                candidate_id,

                exam_id,

                report_type,

                title,

                report_content,

                model_name,
                generated_at

            )

            VALUES(NULL,?,?,?,?,?,?)

        """, (

            exam_id,

            "exam",

            f"Exam Report - {exam['title']}",

            report,

            "ollama",
            generated_at

        ))


        conn.commit()


        # ============================================================
        # 11. RESPONSE
        # ============================================================

        return jsonify({

            "success": True,

            "report_id": cursor.lastrowid,

            "assessment": report,

            "exam": dict(exam),

            "candidate": None,

            "risk": "N/A",

            "integrity_score": 0,

            "face_presence_ratio": 0,

            "warning_count": 0,

            "violations": violations

        })


    except Exception as e:

        conn.rollback()

        return jsonify({

            "success": False,

            "error": str(e)

        }), 500


    finally:

        conn.close()

# ==========================================================
# AI REPORT HISTORY
# ==========================================================

@admin_dashboard_bp.route(
    "/api/ai-reports/history",
    methods=["GET"]
)
@admin_api_required
def ai_report_history():

    conn = get_db_connection()

    try:

        rows = conn.execute("""
            SELECT

                r.id,
                r.candidate_id,
                r.exam_id,
                r.report_type,
                r.title,
                r.risk_label,
                r.integrity_score,
                r.face_presence_ratio,
                r.warning_count,
                r.model_name,
                r.generated_at,

                c.name AS candidate_name,

                e.title AS exam_title

            FROM AIReports r

            LEFT JOIN Candidates c
                ON c.id = r.candidate_id

            INNER JOIN Exams e
                ON e.id = r.exam_id

            ORDER BY
                datetime(r.generated_at) DESC

            LIMIT 100
        """).fetchall()

        reports = []

        for row in rows:

            reports.append({

                "id":
                    row["id"],

                "candidate_id":
                    row["candidate_id"],

                "exam_id":
                    row["exam_id"],

                "report_type":
                    row["report_type"],

                "title":
                    row["title"],

                "candidate_name":
                    row["candidate_name"],

                "exam_title":
                    row["exam_title"],

                "risk_label":
                    row["risk_label"],

                "integrity_score":
                    row["integrity_score"],

                "face_presence_ratio":
                    row["face_presence_ratio"],

                "warning_count":
                    row["warning_count"],

                "model_name":
                    row["model_name"],

                "generated_at":
                    row["generated_at"]

            })

        return jsonify({

            "success": True,

            "reports":
                reports

        })

    except Exception as error:

        return jsonify({

            "success": False,

            "error":
                str(error)

        }), 500

    finally:

        conn.close()

# ==========================================================
# GET GENERATED REPORT
# ==========================================================

@admin_dashboard_bp.route(
    "/api/ai-reports/<int:report_id>",
    methods=["GET"]
)
@admin_required
def get_ai_report(report_id):

    conn = get_db_connection()

    try:

        row = conn.execute("""
            SELECT

                r.*,

                c.name AS candidate_name,

                c.email AS candidate_email,

                e.title AS exam_title,

                e.topic AS exam_topic

            FROM AIReports r

            LEFT JOIN Candidates c
                ON c.id = r.candidate_id

            INNER JOIN Exams e
                ON e.id = r.exam_id

            WHERE r.id = ?

        """, (
            report_id,
        )).fetchone()

        if row is None:

            return jsonify({

                "success": False,

                "error":
                    "Report not found"

            }), 404

        return jsonify({

            "success": True,

            "report": dict(row)

        })

    except Exception as error:

        return jsonify({

            "success": False,

            "error":
                str(error)

        }), 500

    finally:

        conn.close()

# ==========================================================
# DOWNLOAD AI REPORT AS PDF
# ==========================================================

@admin_dashboard_bp.route(
    "/api/ai-reports/<int:report_id>/pdf",
    methods=["GET"]
)
@admin_required
def download_ai_report_pdf(report_id):

    conn = get_db_connection()

    try:

        row = conn.execute("""
            SELECT

                r.*,

                c.name AS candidate_name,

                c.email AS candidate_email,

                e.title AS exam_title,

                e.topic AS exam_topic

            FROM AIReports r

            LEFT JOIN Candidates c
                ON c.id = r.candidate_id

            INNER JOIN Exams e
                ON e.id = r.exam_id

            WHERE r.id = ?

        """, (
            report_id,
        )).fetchone()

        if row is None:

            return jsonify({

                "success": False,

                "error":
                    "Report not found"

            }), 404

        # --------------------------------------------------
        # PDF BUFFER
        # --------------------------------------------------

        buffer = io.BytesIO()

        document = SimpleDocTemplate(

            buffer,

            pagesize=A4,

            rightMargin=45,

            leftMargin=45,

            topMargin=45,

            bottomMargin=45

        )

        styles = getSampleStyleSheet()

        title_style = styles[
            "Title"
        ]

        heading_style = styles[
            "Heading2"
        ]

        body_style = styles[
            "BodyText"
        ]

        title_style.alignment = (
            TA_CENTER
        )

        elements = []

        # --------------------------------------------------
        # TITLE
        # --------------------------------------------------

        elements.append(
            Paragraph(
                "EXAMGUARD AI",
                title_style
            )
        )

        elements.append(
            Spacer(1, 12)
        )

        elements.append(
            Paragraph(
                "AI Integrity Report",
                heading_style
            )
        )

        elements.append(
            Spacer(1, 10)
        )

        # --------------------------------------------------
        # DETAILS
        # --------------------------------------------------

        elements.append(
            Paragraph(
                f"<b>Candidate:</b> "
                f"{row['candidate_name'] or 'N/A'}",
                body_style
            )
        )

        elements.append(
            Paragraph(
                f"<b>Email:</b> "
                f"{row['candidate_email'] or 'N/A'}",
                body_style
            )
        )

        elements.append(
            Paragraph(
                f"<b>Exam:</b> "
                f"{row['exam_title']}",
                body_style
            )
        )

        elements.append(
            Paragraph(
                f"<b>Risk:</b> "
                f"{row['risk_label'] or 'Unknown'}",
                body_style
            )
        )

        elements.append(
            Paragraph(
                f"<b>Integrity Score:</b> "
                f"{row['integrity_score'] or 0}",
                body_style
            )
        )

        elements.append(
            Paragraph(
                f"<b>Face Presence:</b> "
                f"{row['face_presence_ratio'] or 0}",
                body_style
            )
        )

        elements.append(
            Paragraph(
                f"<b>Warnings:</b> "
                f"{row['warning_count'] or 0}",
                body_style
            )
        )

        elements.append(
            Spacer(1, 20)
        )

        # --------------------------------------------------
        # AI REPORT
        # --------------------------------------------------

        report_content = (
            row["report_content"]
            or ""
        )

        for line in report_content.splitlines():

            line = line.strip()

            if not line:
                elements.append(
                    Spacer(1, 8)
                )
                continue

            escaped = (
                line
                .replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
            )

            if (
                line.isupper()
                and len(line) < 60
            ):

                elements.append(
                    Paragraph(
                        f"<b>{escaped}</b>",
                        heading_style
                    )
                )

            else:

                elements.append(
                    Paragraph(
                        escaped,
                        body_style
                    )
                )

            elements.append(
                Spacer(1, 6)
            )

        # --------------------------------------------------
        # FOOTER INFORMATION
        # --------------------------------------------------

        elements.append(
            Spacer(1, 20)
        )

        elements.append(
            Paragraph(
                f"Generated: "
                f"{row['generated_at']}",
                body_style
            )
        )

        elements.append(
            Paragraph(
                f"Model: "
                f"{row['model_name'] or 'Ollama'}",
                body_style
            )
        )

        document.build(
            elements
        )

        buffer.seek(0)

        safe_candidate = (
            row["candidate_name"]
            or "candidate"
        )

        safe_candidate = (
            safe_candidate
            .replace(" ", "_")
            .replace("/", "_")
            .replace("\\", "_")
        )

        filename = (
            f"AI_Integrity_Report_"
            f"{safe_candidate}_"
            f"{report_id}.pdf"
        )

        return send_file(

            buffer,

            mimetype="application/pdf",

            as_attachment=True,

            download_name=filename

        )

    except Exception as error:

        print(
            "PDF generation error:",
            repr(error)
        )

        return jsonify({

            "success": False,

            "error":
                str(error)

        }), 500

    finally:

        conn.close()


@admin_dashboard_bp.route(
    "/api/ai-reports/<int:report_id>/docx",
    methods=["GET"]
)
@admin_required
def download_ai_report_docx(report_id):

    conn = get_db_connection()

    try:

        row = conn.execute("""
            SELECT

                r.*,

                c.name AS candidate_name,

                c.email AS candidate_email,

                e.title AS exam_title

            FROM AIReports r

            LEFT JOIN Candidates c
                ON c.id = r.candidate_id

            INNER JOIN Exams e
                ON e.id = r.exam_id

            WHERE r.id = ?

        """, (report_id,)).fetchone()

        if row is None:

            return jsonify({
                "success": False,
                "error": "Report not found"
            }), 404

        document = Document()

        document.add_heading(
            "ExamGuard AI Report",
            level=1
        )

        document.add_paragraph(
            f"Candidate : {row['candidate_name']}"
        )

        document.add_paragraph(
            f"Email : {row['candidate_email']}"
        )

        document.add_paragraph(
            f"Exam : {row['exam_title']}"
        )

        document.add_paragraph(
            f"Risk : {row['risk_label']}"
        )

        document.add_paragraph(
            f"Integrity Score : {row['integrity_score']}"
        )

        document.add_paragraph(
            f"Warnings : {row['warning_count']}"
        )

        document.add_heading(
            "AI Assessment",
            level=2
        )

        document.add_paragraph(
            row["report_content"] or ""
        )

        buffer = io.BytesIO()

        document.save(buffer)

        buffer.seek(0)

        return send_file(

            buffer,

            as_attachment=True,

            download_name=f"AI_Report_{report_id}.docx",

            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        )

    finally:

        conn.close()

@admin_dashboard_bp.route(
    "/api/ai-reports/<int:report_id>/csv",
    methods=["GET"]
)
@admin_required
def download_ai_report_csv(report_id):

    conn = get_db_connection()

    try:

        row = conn.execute("""
            SELECT

                r.*,

                c.name AS candidate_name,

                c.email AS candidate_email,

                e.title AS exam_title

            FROM AIReports r

            LEFT JOIN Candidates c
                ON c.id = r.candidate_id

            INNER JOIN Exams e
                ON e.id = r.exam_id

            WHERE r.id = ?

        """, (report_id,)).fetchone()

        if row is None:

            return jsonify({
                "success": False,
                "error": "Report not found"
            }), 404

        output = io.StringIO()

        writer = csv.writer(output)

        writer.writerow(["Field", "Value"])

        writer.writerow(["Candidate", row["candidate_name"]])

        writer.writerow(["Email", row["candidate_email"]])

        writer.writerow(["Exam", row["exam_title"]])

        writer.writerow(["Risk", row["risk_label"]])

        writer.writerow(["Integrity Score", row["integrity_score"]])

        writer.writerow(["Warnings", row["warning_count"]])

        writer.writerow(["Assessment", row["report_content"]])

        memory = io.BytesIO()

        memory.write(output.getvalue().encode())

        memory.seek(0)

        return send_file(

            memory,

            as_attachment=True,

            download_name=f"AI_Report_{report_id}.csv",

            mimetype="text/csv"

        )

    finally:

        conn.close()

# ==========================================================
# GET CANDIDATES FOR EXAM (AI REPORTS)
# ==========================================================

@admin_dashboard_bp.route(
    "/api/ai-reports/candidates",
    methods=["GET"]
)
@admin_required
def ai_reports_candidates():

    exam_id = request.args.get(
        "exam_id",
        type=int
    )

    if exam_id is None:

        return jsonify({
            "success": False,
            "error": "exam_id is required"
        }), 400

    conn = get_db_connection()

    try:

        rows = conn.execute(
            """
            SELECT DISTINCT
                c.id,
                c.name,
                c.email
            FROM Candidates c
            WHERE c.id IN (
                SELECT candidate_id FROM SessionLogs WHERE exam_id = ?
                UNION
                SELECT candidate_id FROM ExamAttempts WHERE exam_id = ?
                UNION
                SELECT candidate_id FROM IntegrityScores WHERE exam_id = ?
                UNION
                SELECT candidate_id FROM ViolationLogs WHERE exam_id = ?
            )
            ORDER BY c.name COLLATE NOCASE ASC
            """,
            (exam_id, exam_id, exam_id, exam_id)
        ).fetchall()

        if not rows:
            rows = conn.execute(
                """
                SELECT id, name, email
                FROM Candidates
                ORDER BY name COLLATE NOCASE ASC
                """
            ).fetchall()

        candidates = [
            {
                "id": row["id"],
                "name": row["name"],
                "email": row["email"]
            }
            for row in rows
        ]

        return jsonify({
            "success": True,
            "candidates": candidates
        })

    except Exception as error:

        print(
            "AI report candidates error:",
            repr(error)
        )

        return jsonify({
            "success": False,
            "error": str(error)
        }), 500

    finally:

        conn.close()