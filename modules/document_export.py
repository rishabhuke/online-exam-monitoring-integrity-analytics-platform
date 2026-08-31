"""
PDF/DOCX report rendering (Milestone 5 - integrity analysis port).

Pure rendering layer - takes the exact payload dict already built by
routes.export._build_export_payload() (no new data fetching, no DB access
here) and produces formatted document bytes. Mirrors the section structure
of routes.export._build_export_csv(): session info, integrity score, face
absence events, browser events, flags, AI summary, cluster assignment.

Kept independent of Flask - both functions take a plain dict and return
bytes, so they're callable/testable without a request context.
"""

import io

from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors


def _rows_to_table_data(rows: list) -> list:
    """Converts a list of dicts into a [header_row, *data_rows] structure
    for use in either a reportlab Table or a docx table. Returns an empty
    list if rows is empty (caller decides how to render 'no data')."""
    if not rows:
        return []
    headers = list(rows[0].keys())
    data = [headers]
    for row in rows:
        data.append([str(row.get(h, "")) for h in headers])
    return data


def build_pdf_report(payload: dict) -> bytes:
    """Renders the export payload as a PDF, returned as raw bytes."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    elements = []

    elements.append(Paragraph("Exam Session Integrity Report", styles["Title"]))
    elements.append(Spacer(1, 12))
    elements.append(Paragraph(
        f"Candidate ID: {payload['candidate_id']} &nbsp;&nbsp; Exam ID: {payload['exam_id']}",
        styles["Normal"],
    ))
    elements.append(Spacer(1, 12))

    def add_section(title: str, rows: list):
        elements.append(Paragraph(title, styles["Heading2"]))
        table_data = _rows_to_table_data(rows)
        if not table_data:
            elements.append(Paragraph("(no data)", styles["Normal"]))
        else:
            table = Table(table_data, hAlign="LEFT")
            table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#333333")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]))
            elements.append(table)
        elements.append(Spacer(1, 16))

    add_section("Integrity Score", [payload["integrity_score"]])
    add_section("Face Absence Events", payload["face_absence_events"])
    add_section("Browser Events", payload["browser_events"])
    add_section("Flags", payload["flags"])
    add_section("AI Summary", [payload["ai_summary"]])

    cluster_assignment = payload["cluster_assignment"]
    add_section(
        "Cluster Assignment",
        [cluster_assignment] if cluster_assignment is not None else [],
    )

    doc.build(elements)
    return buffer.getvalue()


def build_docx_report(payload: dict) -> bytes:
    """Renders the export payload as a DOCX, returned as raw bytes."""
    document = Document()

    document.add_heading("Exam Session Integrity Report", level=1)
    session_para = document.add_paragraph()
    session_para.add_run(f"Candidate ID: {payload['candidate_id']}    ").bold = True
    session_para.add_run(f"Exam ID: {payload['exam_id']}").bold = True

    def add_section(title: str, rows: list):
        document.add_heading(title, level=2)
        table_data = _rows_to_table_data(rows)
        if not table_data:
            document.add_paragraph("(no data)")
            return
        headers, *data_rows = table_data
        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Light Grid Accent 1"
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for run in cell.paragraphs[0].runs:
                run.font.bold = True
                run.font.size = Pt(9)
        for row in data_rows:
            cells = table.add_row().cells
            for i, value in enumerate(row):
                cells[i].text = value
                for run in cells[i].paragraphs[0].runs:
                    run.font.size = Pt(9)

    add_section("Integrity Score", [payload["integrity_score"]])
    add_section("Face Absence Events", payload["face_absence_events"])
    add_section("Browser Events", payload["browser_events"])
    add_section("Flags", payload["flags"])
    add_section("AI Summary", [payload["ai_summary"]])

    cluster_assignment = payload["cluster_assignment"]
    add_section(
        "Cluster Assignment",
        [cluster_assignment] if cluster_assignment is not None else [],
    )

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
